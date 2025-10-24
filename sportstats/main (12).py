import sys
import psycopg2
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QComboBox, QDateEdit, QPushButton, QTableView,
    QVBoxLayout, QLabel, QLineEdit, QHBoxLayout, QDialog, QMessageBox, QCheckBox,
    QDialog, QComboBox, QDateEdit, QPushButton, QTableView,
    QMessageBox, QFileDialog, QTextEdit
)
from PyQt6.QtGui import QStandardItemModel, QStandardItem
from PyQt6.uic import loadUi
from collections.abc import Iterable
import re
from docx import Document
from docx.shared import Inches
import os
from PyQt6.QtWidgets import QFileDialog
from PyQt6.QtWidgets import (
    QDialog, QComboBox, QDateEdit, QPushButton, QTableView,
    QMessageBox, QFileDialog
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QStandardItemModel, QStandardItem
from PyQt6.uic import loadUi
from docx import Document
import sys
import os

def resource_path(relative_path):
    """Возвращает абсолютный путь к ресурсу, работает и в .exe и в IDE"""
    if getattr(sys, 'frozen', False):
        # Программа собрана в .exe
        base_path = sys._MEIPASS
    else:
        # Обычный режим (в IDE)
        base_path = os.path.abspath("../../Downloads/Telegram Desktop")
    return os.path.join(base_path, relative_path)


class LoginWindow(QDialog):
    def __init__(self):
        super().__init__()
        self.registration_window = None
        self.conn = None
        self.cursor = None

        loadUi(resource_path("enter.ui"), self)
        self.setFixedSize(self.size())

        self.pushButton_2 = self.findChild(QPushButton, "pushButton_2")
        self.pushButton_3 = self.findChild(QPushButton, "pushButton_3")
        self.lineEdit = self.findChild(QLineEdit, "lineEdit")
        self.lineEdit_2 = self.findChild(QLineEdit, "lineEdit_2")


        self.pushButton_2.clicked.connect(self.connect_to_db)
        self.pushButton_3.clicked.connect(self.open_register_window) # ← вход по кнопке

    @staticmethod
    def create_database_if_not_exists():
        try:
            print("Подключение к системной базе данных для создания...")  # Лог
            with psycopg2.connect(
                    dbname='postgres',  # ← подключаемся к системной БД
                    user='project_admin',
                    password='progectadmin',
                    host='localhost',
                    port=5432,
                    options='-c client_encoding=UTF8'
            ) as conn:
                conn.autocommit = True
                cursor = conn.cursor()
                cursor.execute("SELECT 1 FROM pg_database WHERE datname='sportsstatsdb';")
                exists = cursor.fetchone()
                if not exists:
                    cursor.execute("CREATE DATABASE sportsstatsdb;")
                    print("База данных 'sportsstatsdb' создана.")
                else:
                    print("База данных уже существует.")
        except Exception as e:
            print(f"Ошибка при создании базы данных: {e}")

    def initialize_database_structure(self):
        try:
            with psycopg2.connect(
                    dbname='sportsstatsdb',
                    user='project_admin',
                    password='progectadmin',
                    host='localhost',
                    port=5432,
                    options='-c client_encoding=UTF8'
            ) as conn:
                conn.autocommit = True
                cursor = conn.cursor()
                # Создание ролей, если не существуют
                cursor.execute(
                    "DO $$ BEGIN IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'project_admin') THEN CREATE ROLE project_admin LOGIN PASSWORD 'progectadmin'; END IF; END $$;")
                cursor.execute(
                    "DO $$ BEGIN IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'project_user') THEN CREATE ROLE project_user LOGIN PASSWORD 'projectuser'; END IF; END $$;")

                cursor.execute("""
                -- Создание роли для администратора базы данных
                DO $$ BEGIN
                IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'data_base_admin') THEN
                    CREATE ROLE data_base_admin WITH LOGIN PASSWORD 'admin2';
                END IF;
                END $$;

                GRANT ALL PRIVILEGES ON DATABASE Sportsstatsdb TO data_base_admin;
                """)

                cursor.execute("""
                -- Создание роли для пользователя, работающего с данными только для себя
                DO $$ BEGIN
                IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'user_for_yourself') THEN
                    CREATE ROLE user_for_yourself WITH LOGIN PASSWORD 'baseuser';
                END IF;
                END $$;

                GRANT SELECT, UPDATE, DELETE, INSERT ON ALL TABLES IN SCHEMA public TO user_for_yourself;
                ALTER DEFAULT PRIVILEGES IN SCHEMA public GRANT SELECT, UPDATE, DELETE, INSERT ON TABLES TO user_for_yourself;
                """)

                cursor.execute("""
                -- Создание роли администратора проекта
                DO $$ BEGIN
                IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'project_admin') THEN
                    CREATE ROLE project_admin WITH LOGIN PASSWORD 'progectadmin';
                END IF;
                END $$;

                GRANT SELECT, UPDATE, DELETE, INSERT ON ALL TABLES IN SCHEMA public TO project_admin;
                -- И аналогично далее:
                -- GRANT SELECT, UPDATE, DELETE, INSERT ON SportType TO project_admin;
                -- ...
                """)

                # Таблица пользователей
                cursor.execute("""
                                    -- Creating the SportType table
                CREATE TABLE IF NOT EXISTS public.SportType (
                    ID_SportType INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    SportName VARCHAR(255) NOT NULL,
                    Description TEXT,
                    Rules TEXT
                );

                -- Creating the Team table
                CREATE TABLE IF NOT EXISTS public.Team (
                    ID_Team INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    TeamName VARCHAR(255) NOT NULL,
                    Coach VARCHAR(255),
                    Location VARCHAR(255),
                    ID_SportType INT,
                    FOREIGN KEY (ID_SportType) REFERENCES SportType(ID_SportType) ON DELETE SET NULL
                );

                -- Creating the Player table
                CREATE TABLE IF NOT EXISTS public.Player (
                    ID_Player INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    FirstName VARCHAR(255) NOT NULL,
                    LastName VARCHAR(255) NOT NULL,
                    DateOfBirth DATE,
                    Position VARCHAR(255), -- e.g., "Forward", "Defender"
                    ID_Team INT,
                    FOREIGN KEY (ID_Team) REFERENCES Team(ID_Team) ON DELETE SET NULL
                );

                -- Creating the Tournament table
                CREATE TABLE IF NOT EXISTS public.Tournament (
                    ID_Tournament INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    TournamentName VARCHAR(255) NOT NULL,
                    StartDate DATE,
                    EndDate DATE,
                    ID_SportType INT,
                    FOREIGN KEY (ID_SportType) REFERENCES SportType(ID_SportType) ON DELETE SET NULL
                );

                -- Creating the Match table
                CREATE TABLE IF NOT EXISTS public.Match (
                    ID_Match INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    ID_Tournament INT,
                    ID_SportType INT,
                    MatchDateTime TIMESTAMP WITHOUT TIME ZONE,
                    Location VARCHAR(255),
                    ID_Team1 INT NOT NULL,
                    ID_Team2 INT NOT NULL,
                    FOREIGN KEY (ID_Tournament) REFERENCES Tournament(ID_Tournament) ON DELETE SET NULL,
                    FOREIGN KEY (ID_SportType) REFERENCES SportType(ID_SportType) ON DELETE SET NULL,
                    FOREIGN KEY (ID_Team1) REFERENCES Team(ID_Team),
                    FOREIGN KEY (ID_Team2) REFERENCES Team(ID_Team)
                );

                -- Creating the Result table
                CREATE TABLE IF NOT EXISTS public.Result (
                    ID_Result INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    ID_Match INT NOT NULL,
                    ID_Team INT NOT NULL,
                    Score INT,  -- Total score of the team in the match
                    FOREIGN KEY (ID_Match) REFERENCES Match(ID_Match) ON DELETE CASCADE,
                    FOREIGN KEY (ID_Team) REFERENCES Team(ID_Team)
                );

                -- Creating the PlayerStats table
                CREATE TABLE IF NOT EXISTS public.PlayerStats (
                    ID_Stats INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    ID_Player INT NOT NULL,
                    ID_Match INT NOT NULL,
                    StatType VARCHAR(255),  -- e.g., "Goals", "Assists", "Shots on Goal"
                    Value INT,
                    FOREIGN KEY (ID_Player) REFERENCES Player(ID_Player) ON DELETE CASCADE,
                    FOREIGN KEY (ID_Match) REFERENCES Match(ID_Match) ON DELETE CASCADE
                );
                """)

                cursor.execute("""
                CREATE TABLE IF NOT EXISTS dbusers (
                user_email VARCHAR(255) PRIMARY KEY,
                user_password VARCHAR(255) NOT NULL
                );
                """)

                # Тестовый пользователь
                cursor.execute("""
                    INSERT INTO dbusers (user_email, user_password)
                    VALUES ('test@example.com', '12345')
                    ON CONFLICT (user_email) DO NOTHING;
                """)

                print("Структура БД и роли инициализированы.")
        except Exception as e:
            print("Ошибка инициализации БД:", e)

    def connect_to_db(self):
        self.create_database_if_not_exists()
        # Получение текста из lineEdit
        user_input = self.lineEdit.text()
        password_input = self.lineEdit_2.text()

        try:
            # Проверяем, если это данные администратора
            if user_input == "project_admin" and password_input == "progectadmin":
                # Подключение как администратор
                self.user_role = "project_admin"

                self.conn = psycopg2.connect(
                    dbname='sportsstatsdb',
                    user='project_admin',
                    password='progectadmin',
                    host='localhost',
                    port=5432,
                    options='-c client_encoding=UTF8'
                )
                self.conn.autocommit = True
                self.cursor = self.conn.cursor()

                self.open_user_window()  # Открываем окно администратора
                return

            # Если это не администратор, выполняем поиск в базе данных
            with psycopg2.connect(
                    dbname='sportsstatsdb',
                    user='project_admin',
                    password='progectadmin',
                    host='localhost',
                    port=5432,
                    options='-c client_encoding=UTF8'
            ) as admin_conn:
                admin_cursor = admin_conn.cursor()
                admin_cursor.execute(
                    "SELECT user_password FROM dbusers WHERE user_email = %s",
                    (user_input,)
                )
                result = admin_cursor.fetchone()

                if result is not None:
                    stored_password = str(result[0])  # Преобразуем в строку
                    print(f"Сохранённый пароль: {stored_password}")  # Лог

                    # Если пароль совпадает, подключаемся как обычный пользователь
                    if password_input == stored_password:
                        print("Пароль верен")
                        self.user_role = "project_user"

                        # Подключение как обычный пользователь
                        self.conn = psycopg2.connect(
                            dbname='sportsstatsdb',
                            user='project_user',
                            password='projectuser',
                            host='localhost',
                            port=5432,
                            options='-c client_encoding=UTF8'
                        )
                        self.conn.autocommit = True
                        self.cursor = self.conn.cursor()

                        self.open_user_window()  # Открываем окно пользователя

                    else:
                        # Если пароль не совпадает, показываем сообщение об ошибке
                        QMessageBox.warning(self, "Ошибка", "Неверный пароль.")
                else:
                    # Если пользователь не найден, показываем сообщение об ошибке
                    QMessageBox.warning(self, "Ошибка", "Пользователь не найден.")

        except Exception as e:
            # В случае ошибки при подключении, показываем сообщение с ошибкой
            QMessageBox.critical(self, "Ошибка", str(e))

    def execute_initial_sql(self):
        sql_statements = [
            """
            CREATE OR REPLACE FUNCTION create_match_result()
            RETURNS TRIGGER AS $$
            BEGIN
                INSERT INTO Result (ID_Match, ID_Team, Score)
                VALUES (NEW.ID_Match, NEW.ID_Team1, 0),
                       (NEW.ID_Match, NEW.ID_Team2, 0);
                RETURN NEW;
            END;
            $$ LANGUAGE plpgsql;
            """,
            """
            CREATE OR REPLACE FUNCTION update_goals_stats()
            RETURNS TRIGGER AS $$
            BEGIN
                IF NEW.StatType = 'Goals' THEN
                    UPDATE PlayerStats
                    SET Value = Value + NEW.Value
                    WHERE ID_Player = NEW.ID_Player AND ID_Match = NEW.ID_Match AND StatType = 'Goals';
                END IF;
                RETURN NEW;
            END;
            $$ LANGUAGE plpgsql;

            """,
            """
            CREATE OR REPLACE FUNCTION prevent_tournament_deletion()
            RETURNS TRIGGER AS $$
            BEGIN
                IF EXISTS (SELECT 1 FROM Match WHERE ID_Tournament = OLD.ID_Tournament) THEN
                    RAISE EXCEPTION 'Нельзя удалить турнир с проведенными матчами';
                END IF;
                RETURN OLD;
            END;
            $$ LANGUAGE plpgsql;
            """,
            """
            CREATE VIEW match_summary_view AS
            SELECT m.ID_Match, m.MatchDateTime, m.Location, t1.TeamName AS Team1, t2.TeamName AS Team2
            FROM Match m
            JOIN Team t1 ON m.ID_Team1 = t1.ID_Team
            JOIN Team t2 ON m.ID_Team2 = t2.ID_Team;
            """,
            """
                CREATE TRIGGER trg_create_match_result
                AFTER INSERT ON Match
                FOR EACH ROW
                EXECUTE FUNCTION create_match_result();
                """,
            """
            CREATE TRIGGER trg_prevent_tournament_deletion
            BEFORE DELETE ON Tournament
            FOR EACH ROW
            EXECUTE FUNCTION prevent_tournament_deletion();
            """
        ]

        for sql in sql_statements:
            try:
                self.cursor.execute(sql)
                print("SQL-команда выполнена успешно.")
            except psycopg2.Error as e:
                print("Ошибка SQL:", e)

        try:
            self.cursor.execute("BEGIN;")
            self.cursor.execute("SAVEPOINT sp1;")
            self.cursor.execute("SAVEPOINT sp2;")
            self.cursor.execute("ROLLBACK TO SAVEPOINT sp1;")
            self.cursor.execute("COMMIT;")
        except Exception as e:
            self.cursor.execute("ROLLBACK;")
            print("Ошибка точек сохранения:", e)

    def open_register_window(self):
        self.user_window = RegisterWindow(self.conn, self.cursor)
        self.user_window.show()

    def open_user_window(self):
        self.user_window = UserWindow(self.conn, self.cursor, self.user_role)
        self.user_window.show()
        self.hide()


class RegisterWindow(QDialog):
    def __init__(self, conn, cursor):
        super().__init__()
        loadUi(resource_path("registration.ui"), self)

        # Привязка виджетов
        self.lineEdit = self.findChild(QLineEdit, "lineEdit")  # Почта
        self.lineEdit_2 = self.findChild(QLineEdit, "lineEdit_2")  # Пароль
        self.lineEdit_3 = self.findChild(QLineEdit, "lineEdit_3")  # Повтор пароля
        self.pushButton_2 = self.findChild(QPushButton, "pushButton_2")  # Кнопка регистрации

        self.setFixedSize(self.size())
        self.pushButton_2.clicked.connect(self.register_user)

    def register_user(self):
        email = self.lineEdit.text()  # Почта из поля
        password = self.lineEdit_2.text()  # Пароль из поля
        confirm_password = self.lineEdit_3.text()  # Подтверждение пароля из поля

        # Проверка на пустые поля
        if not email or not password or not confirm_password:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все поля.")
            return

        # Проверка корректности почты
        email_regex = r'^[\w\.-]+@[\w\.-]+\.\w+$'
        if not re.match(email_regex, email):
            QMessageBox.warning(self, "Ошибка", "Некорректный адрес электронной почты.")
            return

        # Проверка на совпадение пароля и подтверждения
        if password != confirm_password:
            QMessageBox.warning(self, "Ошибка", "Пароли не совпадают.")
            return

        try:
            # Подключение к базе данных
            with psycopg2.connect(
                    dbname='sportsstatsdb',
                    user='project_admin',
                    password='progectadmin',
                    host='localhost',
                    port=5432,
                    options='-c client_encoding=UTF8'
            ) as admin_conn:
                admin_cursor = admin_conn.cursor()

                # Проверка, существует ли уже пользователь с таким email
                admin_cursor.execute("SELECT 1 FROM dbusers WHERE user_email = %s", (email,))
                if admin_cursor.fetchone():
                    QMessageBox.warning(self, "Ошибка", "Пользователь уже зарегистрирован.")
                    return

                # Вставка нового пользователя в базу данных
                # Явно убеждаемся, что пароль - строка
                admin_cursor.execute("""
                    INSERT INTO dbusers (user_email, user_password)
                    VALUES (%s, %s)
                """, (email, str(password)))  # Преобразуем пароль в строку

                QMessageBox.information(self, "Успешно", "Пользователь зарегистрирован!")

                # Очистка полей
                self.lineEdit.clear()
                self.lineEdit_2.clear()
                self.lineEdit_3.clear()

                self.close()  # Закрытие окна регистрации

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка регистрации: {str(e)}")


class UserWindow(QDialog):
    def __init__(self, conn, cursor, user_role):
        super().__init__()
        self.conn = conn
        self.cursor = cursor
        self.user_role = user_role

        loadUi(resource_path("../../Downloads/Telegram Desktop/design.ui"), self)
        self.setFixedSize(self.size())

        # Подключение виджетов по именам из UI
        self.sport_combo = self.findChild(QComboBox, "comboBox")            # Вид спорта
        self.date_edit = self.findChild(QDateEdit, "dateEdit")              # Дата
        self.result_button = self.findChild(QPushButton, "pushButton")      # Посмотреть результат
        self.add_match_button = self.findChild(QPushButton, "pushButton_2") # Добавить матч
        self.table_view = self.findChild(QTableView, "tableView")           # Таблица поиска матча
        self.table_last_games = self.findChild(QTableView, "tableView_2")   # Таблица недавние матчи
        self.table_future_games = self.findChild(QTableView, "tableView_3") # Таблица последние матчи

        self.tournament_combo = self.findChild(QComboBox, "comboBox_2")     # Турнир
        self.team_combo = self.findChild(QComboBox, "comboBox_3")           # Команда
        self.play_combo = self.findChild(QComboBox, "comboBox_4")           # Для выбрать игру
        self.tournament_combo_5 = self.findChild(QComboBox, "comboBox_5")   # Для выбрать турнир
        self.label_sport_type = self.findChild(QLabel, "label")
        self.label_team_search = self.findChild(QLabel, "label_4")
        self.label_search_tournament = self.findChild(QLabel, "label_3")
        self.label_select_date = self.findChild(QLabel, "label_2")

        self.checkbox_sport_type = self.findChild(QCheckBox, "checkBox")
        self.checkbox_team = self.findChild(QCheckBox, "checkBox_4")
        self.checkbox_tournament = self.findChild(QCheckBox, "checkBox_3")
        self.checkbox_date = self.findChild(QCheckBox, "checkBox_2")

        self.label_sport_type.hide()
        self.tournament_combo.hide()
        self.label_team_search.hide()
        self.team_combo.hide()
        self.label_search_tournament.hide()
        self.sport_combo.hide()
        self.label_select_date.hide()
        self.date_edit.hide()

        self.checkbox_sport_type.stateChanged.connect(self.toggle_sport_visibly)
        self.checkbox_team.stateChanged.connect(self.toggle_team_visibly)
        self.checkbox_tournament.stateChanged.connect(self.toggle_tournament_visibly)
        self.checkbox_date.stateChanged.connect(self.toggle_date_visibly)

        if self.user_role != "project_admin":
            self.add_match_button.hide()

        self.result_button.clicked.connect(self.load_data)
        self.add_match_button.clicked.connect(self.add_match)

        self.load_sport_types()
        self.load_tournaments()
        self.load_teams()
        self.table_view.clicked.connect(self.on_table_row_clicked)
        self.table_teams_names.clicked.connect(self.handle_row_click)

    def toggle_sport_visibly(self, state):
        if state.isChecked:
            self.label_sport_type.show()
            self.tournament_combo.show()
        else:
            self.label_sport_type.hide()
            self.tournament_combo.hide()

    def toggle_team_visibly(self, state):
        if state.isChecked:
            self.label_team_search.show()
            self.team_combo.show()
        else:
            self.label_team_search.hide()
            self.team_combo.hide()

    def toggle_tournament_visibly(self, state):
        if state.isChecked:
            self.label_search_tournament.show()
            self.sport_combo.show()
        else:
            self.label_search_tournament.hide()
            self.sport_combo.hide()

    def toggle_date_visibly(self, state):
        if state.isChecked:
            self.label_select_date.show()
            self.date_edit.show()
        else:
            self.label_select_date.hide()
            self.date_edit.hide()


    def load_sport_types(self):
        try:
            self.cursor.execute("SELECT SportName FROM SportType;")
            sports = self.cursor.fetchall()
            self.sport_combo.clear()
            for sport in sports:
                self.sport_combo.addItem(sport[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки видов спорта:\n{e}")

    def load_tournaments(self):
        try:
            self.cursor.execute("SELECT TournamentName FROM Tournament;")
            tournaments = self.cursor.fetchall()
            self.tournament_combo.clear()
            self.tournament_combo.addItem("Все турниры")
            for tournament in tournaments:
                self.tournament_combo.addItem(tournament[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки турниров:\n{e}")

    def load_teams(self):
        try:
            self.cursor.execute("SELECT TeamName FROM Team;")
            teams = self.cursor.fetchall()
            self.team_combo.clear()
            self.team_combo.addItem("Все команды")
            for team in teams:
                self.team_combo.addItem(team[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки команд:\n{e}")

    def load_data(self):
        selected_sport = self.sport_combo.currentText()
        selected_date = self.date_edit.date().toString("yyyy-MM-dd")
        selected_team = self.team_combo.currentText()
        selected_tournament = self.tournament_combo.currentText()
        selected_play = self.play_combo.currentText()
        selected_tournament5 = self.tournament_combo_5.currentText()

        try:
            query = """
                SELECT m.MatchDateTime, t1.TeamName, t2.TeamName, r1.Score, r2.Score
                FROM Match m
                JOIN Team t1 ON m.ID_Team1 = t1.ID_Team
                JOIN Team t2 ON m.ID_Team2 = t2.ID_Team
                JOIN SportType st ON m.ID_SportType = st.ID_SportType
                LEFT JOIN Result r1 ON r1.ID_Match = m.ID_Match AND r1.ID_Team = t1.ID_Team
                LEFT JOIN Result r2 ON r2.ID_Match = m.ID_Match AND r2.ID_Team = t2.ID_Team
                LEFT JOIN Tournament tr ON m.ID_Tournament = tr.ID_Tournament
                WHERE 1=1
            """

            params = []

            if selected_sport and selected_sport != "Все":
                query += " AND st.SportName = %s"
                params.append(selected_sport)

            if selected_play and selected_play != "Все":
                query += " AND st.SportName = %s"
                params.append(selected_play)

            if selected_tournament and selected_tournament != "Все турниры":
                query += " AND tr.TournamentName = %s"
                params.append(selected_tournament)

            if selected_tournament5 and selected_tournament5 != "Все":
                query += " AND tr.TournamentName = %s"
                params.append(selected_tournament5)

            if selected_team and selected_team != "Все команды":
                query += " AND (t1.TeamName = %s OR t2.TeamName = %s)"
                params.extend([selected_team, selected_team])

            if self.date_edit.date().isValid():
                query += " AND DATE(m.MatchDateTime) = %s"
                params.append(selected_date)

            self.cursor.execute(query, params)
            data = self.cursor.fetchall()

            if data:
                model = QStandardItemModel(self)
                model.setColumnCount(5)
                model.setHorizontalHeaderLabels(["Дата", "Команда 1", "Счёт 1", "Счёт 2", "Команда 2"])

                for row in data:
                    match_date, team1, team2, score1, score2 = row
                    items = [
                        QStandardItem(str(match_date)),
                        QStandardItem(team1),
                        QStandardItem(str(score1) if score1 is not None else "-"),
                        QStandardItem(str(score2) if score2 is not None else "-"),
                        QStandardItem(team2),
                    ]
                    model.appendRow(items)

                self.table_view.setModel(model)

                reply = QMessageBox.question(
                    self,
                    "Сохранить?",
                    "Сохранить результаты в Word документ?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )

                if reply == QMessageBox.StandardButton.Yes:
                    self.export_to_docx(data)
            else:
                if selected_team != "Все команды":
                    self.show_team_players(selected_team)
                else:
                    QMessageBox.information(self, "Информация", "Нет данных для отображения.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки данных:\n{e}")

    def show_team_players(self, team_name):
        try:
            self.cursor.execute("""
                SELECT p.FirstName, p.LastName, p.Position, p.DateOfBirth
                FROM Player p
                JOIN Team t ON p.ID_Team = t.ID_Team
                WHERE t.TeamName = %s;
            """, (team_name,))
            players = self.cursor.fetchall()

            if players:
                model = QStandardItemModel(self)
                model.setColumnCount(4)
                model.setHorizontalHeaderLabels(["Имя", "Фамилия", "Позиция", "Дата рождения"])

                for row in players:
                    items = [QStandardItem(str(cell)) for cell in row]
                    model.appendRow(items)

                self.table_view.setModel(model)
            else:
                QMessageBox.information(self, "Информация", "У этой команды нет игроков.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки игроков:\n{e}")

    def export_to_docx(self, data):
        try:
            doc = Document()
            doc.add_heading('Результаты матчей', 0)

            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Команда 1'
            hdr_cells[2].text = 'Счёт 1'
            hdr_cells[3].text = 'Счёт 2'
            hdr_cells[4].text = 'Команда 2'

            for row in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row[0])
                row_cells[1].text = str(row[1])
                row_cells[2].text = str(row[3]) if row[3] is not None else "-"
                row_cells[3].text = str(row[4]) if row[4] is not None else "-"
                row_cells[4].text = str(row[2])

            # Сохраняем через диалог
            filename, _ = QFileDialog.getSaveFileName(self, "Сохранить как", "", "Word Documents (*.docx)")
            if filename:
                if not filename.endswith(".docx"):
                    filename += ".docx"
                doc.save(filename)
                QMessageBox.information(self, "Успешно", "Документ сохранён.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить документ:\n{e}")

    def add_match(self):
        self.add_result_window = AddResult(self.conn, self.cursor)
        self.add_result_window.show()

    def on_table_row_clicked(self, index):
        if not index.isValid():
            return

        model = self.table_view.model()
        row = index.row()

        match_date = model.item(row, 0).text()
        team1 = model.item(row, 1).text()
        score1 = model.item(row, 2).text()
        score2 = model.item(row, 3).text()
        team2 = model.item(row, 4).text()

        # можно передавать и ID матча, если доступен
        self.match_detail = MatchDetail(self.conn, self.cursor, match_date, team1, team2)
        self.match_detail.exec()

class MatchDetail(QDialog):
    def __init__(self, conn, cursor):
        super().__init__()

        if not hasattr(cursor, "execute"):
            raise ValueError("Ошибка: передан неправильный объект вместо курсора базы данных.")

        self.conn = conn
        self.cursor = cursor
        loadUi(resource_path("details.ui"), self)
        self.setFixedSize(self.size())

        self.table_teams_names = self.findChild(QTableView, "tableView_4")
        self.table_team_one = self.findChild(QTableView, "tableView_5")
        self.table_team_two = self.findChild(QTableView, "tableView_6")
        self.table_detail_one = self.findChild(QTableView, "tableView_7")
        self.table_detail_two = self.findChild(QTableView, "tableView_8")
        self.list_analytic_one = self.findChild(QTextEdit, "textEdit")
        self.list_analytic_two = self.findChild(QTextEdit, "textEdit_2")

        self.list_analytic_one.setReadOnly(True)
        self.list_analytic_two.setReadOnly(True)

        self.load_matches()
        self.table_teams_names.clicked.connect(self.load_players_for_selected_match)

    def load_matches(self):
        try:
            self.cursor.execute("""
                SELECT t1.TeamName, t2.TeamName
                FROM Match m
                JOIN Team t1 ON m.ID_Team1 = t1.ID_Team
                JOIN Team t2 ON m.ID_Team2 = t2.ID_Team
            """)
            rows = self.cursor.fetchall()

            self.match_data = rows  # Сохраняем данные для индексации
            model = QStandardItemModel(self)
            model.setColumnCount(2)
            model.setHorizontalHeaderLabels(["Команда 1", "Команда 2"])

            for team1, team2 in rows:
                model.appendRow([QStandardItem(team1), QStandardItem(team2)])

            self.table_teams_names.setModel(model)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки матчей:\n{e}")

    def load_players_for_selected_match(self, index):
        row = index.row()
        team1_name, team2_name = self.match_data[row]

        self.load_team_players(team1_name, self.table_team_one)
        self.load_team_players(team2_name, self.table_team_two)

        # Можно также загрузить аналитику через OpenAI
        self.generate_analytics(team1_name, self.list_analytic_one)
        self.generate_analytics(team2_name, self.list_analytic_two)

    def load_team_players(self, team_name, table_view):
        try:
            self.cursor.execute("""
                SELECT FirstName, LastName, Position, DateOfBirth
                FROM Player p
                JOIN Team t ON p.ID_Team = t.ID_Team
                WHERE t.TeamName = %s
            """, (team_name,))
            players = self.cursor.fetchall()

            model = QStandardItemModel(self)
            model.setColumnCount(4)
            model.setHorizontalHeaderLabels(["Имя", "Фамилия", "Позиция", "Дата рождения"])

            for row in players:
                items = [QStandardItem(str(cell)) for cell in row]
                model.appendRow(items)

            table_view.setModel(model)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки игроков для {team_name}:\n{e}")

    def generate_analytics(self, team_name, text_edit):
        """
        Безопасная генерация аналитики с защитой от:
        - ошибок сети
        - таймаутов
        - неожиданных структур ответа
        - утечек памяти
        """
        try:
            # Устанавливаем текст ожидания
            text_edit.setText("Идет анализ команды...")

            # Принудительно обновляем интерфейс
            QApplication.processEvents()

            # Импортируем клиент внутри метода для изоляции возможных ошибок импорта
            try:
                from g4f.client import Client
                from g4f.Provider import RetryProvider, Phind, FreeChatgpt
            except ImportError as e:
                text_edit.setText("Ошибка: не удалось импортировать g4f")
                print(f"Import error: {e}")
                return

            # Создаем клиент с таймаутами и повторными попытками
            client = Client(
                provider=RetryProvider([Phind, FreeChatgpt], shuffle=False),
                timeout=30  # 30 секунд на запрос
            )

            # Формируем промпт с ограничением длины
            prompt = (
                f"Проанализируй команду {team_name[:50]}. "
                "Опиши кратко (максимум 100 слов):\n"
                "- Сильные стороны\n"
                "- Слабые стороны\n"
                "- Рекомендации по улучшению"
            )

            # Выполняем запрос с обработкой возможных ошибок
            try:
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",  # Более стабильная модель
                    messages=[{"role": "user", "content": prompt}],
                    web_search=False,
                    stream=False  # Отключаем потоковый вывод для стабильности
                )

                # Безопасное извлечение текста ответа
                if hasattr(response, 'choices'):
                    analysis_text = response.choices[0].message.content
                else:
                    analysis_text = "Не удалось разобрать ответ сервиса"

            except Exception as api_error:
                print(f"API error: {api_error}")
                analysis_text = f"Ошибка запроса: {str(api_error)[:200]}"

            # Ограничиваем длину и устанавливаем текст
            text_edit.setText(analysis_text[:2000])  # Максимум 2000 символов

        except Exception as e:
            error_msg = f"Критическая ошибка: {str(e)[:300]}"
            print(error_msg)
            text_edit.setText(error_msg)
        finally:
            # Принудительное обновление интерфейса
            QApplication.processEvents()


class AddResult(QDialog):
    def __init__(self, conn, cursor):
        super().__init__()
        if not hasattr(cursor, "execute"):
            raise ValueError("Ошибка: передан неправильный объект вместо курсора базы данных.")

        self.conn = conn
        self.cursor = cursor
        loadUi(resource_path("../../Downloads/Telegram Desktop/addmatch.ui"), self)
        self.setFixedSize(self.size())

        self.sport_combo = self.findChild(QComboBox, "comboBox")
        self.tournament_combo = self.findChild(QComboBox, "comboBox_2")
        self.date_edit = self.findChild(QDateEdit, "dateEdit")
        self.team1_line = self.findChild(QLineEdit, "lineEdit")
        self.team2_line = self.findChild(QLineEdit, "lineEdit_2")
        self.score_line = self.findChild(QLineEdit, "lineEdit_3")
        self.add_button = self.findChild(QPushButton, "pushButton_2")

        self.add_button.clicked.connect(self.save_result)
        self.load_sports()
        self.load_tournaments()

    def load_sports(self):
        try:
            self.cursor.execute("SELECT SportName FROM SportType;")
            sports = self.cursor.fetchall()
            self.sport_combo.clear()
            for sport in sports:
                self.sport_combo.addItem(sport[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки видов спорта:\n{e}")

    def load_tournaments(self):
        try:
            self.cursor.execute("SELECT TournamentName FROM Tournament;")
            tournaments = self.cursor.fetchall()
            self.tournament_combo.clear()
            self.tournament_combo.addItem("Без турнира")
            for tournament in tournaments:
                self.tournament_combo.addItem(tournament[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки турниров:\n{e}")

    def save_result(self):
        try:
            data = self.validate_and_get_input()
            if not data:
                return

            sport_name, tournament_name, match_date, team1_name, team2_name, score1, score2 = data

            sport_id = self.get_sport_id(sport_name)
            if sport_id is None:
                QMessageBox.critical(self, "Ошибка", "Вид спорта не найден.")
                return

            tournament_id = self.get_tournament_id(tournament_name) if tournament_name != "Без турнира" else None

            team1_id = self.get_or_create_team(team1_name, sport_id)
            team2_id = self.get_or_create_team(team2_name, sport_id)

            match_id = self.insert_match(match_date, team1_id, team2_id, sport_id, tournament_id)
            self.insert_result(match_id, team1_id, score1, team2_id, score2)

            self.conn.commit()
            QMessageBox.information(self, "Успешно", "Матч и результат добавлены!")

            self.team1_line.clear()
            self.team2_line.clear()
            self.score_line.clear()

        except Exception as e:
            self.conn.rollback()
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить результат:\n{e}")

    def validate_and_get_input(self):
        sport_name = self.sport_combo.currentText()
        tournament_name = self.tournament_combo.currentText()
        match_date = self.date_edit.date().toString("yyyy-MM-dd")
        team1_name = self.team1_line.text().strip()
        team2_name = self.team2_line.text().strip()
        score_text = self.score_line.text().strip()

        if not (team1_name and team2_name and score_text):
            QMessageBox.warning(self, "Проверка", "Пожалуйста, заполните все поля.")
            return None

        scores = score_text.split()
        if len(scores) != 2 or not all(s.isdigit() for s in scores):
            QMessageBox.warning(self, "Неверный формат", "Введите счёт в формате: `1 2`.")
            return None

        return sport_name, tournament_name, match_date, team1_name, team2_name, int(scores[0]), int(scores[1])

    def get_sport_id(self, sport_name):
        self.cursor.execute("SELECT ID_SportType FROM SportType WHERE SportName = %s;", (sport_name,))
        result = self.cursor.fetchone()
        return result[0] if result else None

    def get_tournament_id(self, name):
        self.cursor.execute("SELECT ID_Tournament FROM Tournament WHERE TournamentName = %s;", (name,))
        result = self.cursor.fetchone()
        return result[0] if result else None

    def get_or_create_team(self, team_name, sport_id):
        self.cursor.execute("SELECT ID_Team FROM Team WHERE TeamName = %s;", (team_name,))
        result = self.cursor.fetchone()
        if result:
            return result[0]
        self.cursor.execute(
            "INSERT INTO Team (TeamName, ID_SportType) VALUES (%s, %s) RETURNING ID_Team;",
            (team_name, sport_id)
        )
        return self.cursor.fetchone()[0]

    def insert_match(self, match_date, team1_id, team2_id, sport_id, tournament_id):
        self.cursor.execute("""
            INSERT INTO Match (MatchDateTime, Location, ID_Team1, ID_Team2, ID_SportType, ID_Tournament)
            VALUES (%s, 'Не указано', %s, %s, %s, %s)
            RETURNING ID_Match;
        """, (match_date, team1_id, team2_id, sport_id, tournament_id))
        return self.cursor.fetchone()[0]

    def insert_result(self, match_id, team1_id, score1, team2_id, score2):
        self.cursor.execute("""
            SELECT COUNT(*) FROM Result WHERE ID_Match = %s;
        """, (match_id,))
        already_exists = self.cursor.fetchone()[0] > 0
        if already_exists:
            print(f"Результаты для матча {match_id} уже существуют. Пропуск вставки.")
            return

        self.cursor.execute("""
            INSERT INTO Result (ID_Match, ID_Team, Score)
            VALUES (%s, %s, %s), (%s, %s, %s);
        """, (match_id, team1_id, score1, match_id, team2_id, score2))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = LoginWindow()
    window.show()
    sys.exit(app.exec())
