import traceback
from contextlib import contextmanager
from datetime import time, date
import datetime

import psycopg2
from PyQt6 import uic
from PyQt6.QtCore import QDate, QTimer, Qt, pyqtSlot, QMetaObject
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QComboBox, QDateEdit, QPushButton, QTableView,
    QVBoxLayout, QLabel, QLineEdit, QHBoxLayout, QDialog, QMessageBox, QCheckBox, QTextEdit, QHeaderView,
    QTableWidgetItem, QWidget, QLayout, QTableWidget, QTabWidget, QSpinBox
)
import re
from PyQt6.QtWidgets import (
    QDialog, QComboBox, QDateEdit, QPushButton, QTableView,
    QMessageBox, QFileDialog
)
from PyQt6.uic import loadUi
from docx import Document

from PyQt6.QtCore import QThread, pyqtSignal, QObject, QTimer
from PyQt6.QtWidgets import QDialog, QMessageBox, QHeaderView
from PyQt6.QtGui import QStandardItemModel, QStandardItem
import os
import sys

from g4f import requests
from g4f.client import Client
from g4f.Provider import (
    You,        # Самый стабильный бесплатный
    FreeGpt,    # Альтернатива
    Gemini,     # Google Gemini
    DeepSeek    # Новый перспективный
)

from PyQt6.QtMultimedia import QSoundEffect
from PyQt6.QtCore import QUrl
from functools import wraps
from contextlib import contextmanager



class SoundManager:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance.sounds = {}
        return cls._instance

    def load_sound(self, name, path):
        sound = QSoundEffect()
        sound.setSource(QUrl.fromLocalFile(path))
        self.sounds[name] = sound

    def play(self, name):
        if name in self.sounds:
            self.sounds[name].play()


# Глобальный экземпляр менеджера звуков
sound_manager = SoundManager()
sound_manager.load_sound("click", "click.wav")


def with_sound(sound_name="click"):
    """Декоратор для добавления звукового сопровождения к методам"""

    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            sound_manager.play(sound_name)
            return func(*args, **kwargs)

        return wrapper

    return decorator


def resource_path(relative_path):
    """Возвращает абсолютный путь к ресурсу, работает и в .exe и в IDE"""
    if getattr(sys, 'frozen', False):
        # Программа собрана в .exe
        base_path = sys._MEIPASS
    else:
        # Обычный режим (в IDE)
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


class PharmacyTransactionManager:
    """Менеджер транзакций для аптечной системы"""
    
    def __init__(self, conn, cursor):
        self.conn = conn
        self.cursor = cursor
        self.savepoints = []
        self.transaction_active = False
    
    @contextmanager
    def transaction(self, operation_name="Операция"):
        """Контекстный менеджер для транзакций"""
        try:
            print(f"Начало транзакции: {operation_name}")
            self.cursor.execute("BEGIN;")
            self.transaction_active = True
            yield self
            self.cursor.execute("COMMIT;")
            self.transaction_active = False
            print(f"Транзакция завершена успешно: {operation_name}")
        except Exception as e:
            self.cursor.execute("ROLLBACK;")
            self.transaction_active = False
            print(f"Ошибка транзакции {operation_name}: {e}")
            raise
    
    def create_savepoint(self, name):
        """Создание точки сохранения"""
        try:
            savepoint_name = f"sp_{name}_{len(self.savepoints)}"
            self.cursor.execute(f"SAVEPOINT {savepoint_name};")
            self.savepoints.append(savepoint_name)
            print(f"Создана точка сохранения: {savepoint_name}")
            return savepoint_name
        except Exception as e:
            print(f"Ошибка создания точки сохранения: {e}")
            raise
    
    def rollback_to_savepoint(self, savepoint_name):
        """Откат к точке сохранения"""
        try:
            if savepoint_name in self.savepoints:
                self.cursor.execute(f"ROLLBACK TO SAVEPOINT {savepoint_name};")
                # Удаляем все точки сохранения после этой
                index = self.savepoints.index(savepoint_name)
                self.savepoints = self.savepoints[:index + 1]
                print(f"Откат к точке сохранения: {savepoint_name}")
            else:
                raise ValueError(f"Точка сохранения {savepoint_name} не найдена")
        except Exception as e:
            print(f"Ошибка отката к точке сохранения: {e}")
            raise
    
    def release_savepoint(self, savepoint_name):
        """Освобождение точки сохранения"""
        try:
            if savepoint_name in self.savepoints:
                self.cursor.execute(f"RELEASE SAVEPOINT {savepoint_name};")
                self.savepoints.remove(savepoint_name)
                print(f"Освобождена точка сохранения: {savepoint_name}")
        except Exception as e:
            print(f"Ошибка освобождения точки сохранения: {e}")
    
    def rollback_all(self):
        """Откат всей транзакции"""
        try:
            self.cursor.execute("ROLLBACK;")
            self.savepoints.clear()
            self.transaction_active = False
            print("Откат всей транзакции")
        except Exception as e:
            print(f"Ошибка отката транзакции: {e}")
    
    def commit_all(self):
        """Подтверждение всей транзакции"""
        try:
            self.cursor.execute("COMMIT;")
            self.savepoints.clear()
            self.transaction_active = False
            print("Подтверждение всей транзакции")
        except Exception as e:
            print(f"Ошибка подтверждения транзакции: {e}")
            raise


class PharmacySavePoints:
    """Класс для управления точками сохранения аптечных операций"""
    
    # Константы для точек сохранения
    SALE_START = "sale_start"
    SALE_CUSTOMER_DATA = "sale_customer_data"
    SALE_ITEMS_ADDED = "sale_items_added"
    SALE_INVENTORY_UPDATED = "sale_inventory_updated"
    SALE_COMPLETED = "sale_completed"
    
    INVENTORY_START = "inventory_start"
    INVENTORY_MEDICINE_CHECKED = "inventory_medicine_checked"
    INVENTORY_QUANTITY_UPDATED = "inventory_quantity_updated"
    INVENTORY_COMPLETED = "inventory_completed"
    
    MEDICINE_START = "medicine_start"
    MEDICINE_VALIDATED = "medicine_validated"
    MEDICINE_SUPPLIER_CHECKED = "medicine_supplier_checked"
    MEDICINE_COMPLETED = "medicine_completed"
    
    SUPPLIER_START = "supplier_start"
    SUPPLIER_VALIDATED = "supplier_validated"
    SUPPLIER_COMPLETED = "supplier_completed"
    
    @staticmethod
    def get_savepoint_description(savepoint_name):
        """Получение описания точки сохранения"""
        descriptions = {
            PharmacySavePoints.SALE_START: "Начало операции продажи",
            PharmacySavePoints.SALE_CUSTOMER_DATA: "Данные покупателя сохранены",
            PharmacySavePoints.SALE_ITEMS_ADDED: "Товары добавлены в продажу",
            PharmacySavePoints.SALE_INVENTORY_UPDATED: "Остатки на складе обновлены",
            PharmacySavePoints.SALE_COMPLETED: "Продажа завершена",
            
            PharmacySavePoints.INVENTORY_START: "Начало операции с инвентарем",
            PharmacySavePoints.INVENTORY_MEDICINE_CHECKED: "Лекарство проверено",
            PharmacySavePoints.INVENTORY_QUANTITY_UPDATED: "Количество обновлено",
            PharmacySavePoints.INVENTORY_COMPLETED: "Операция с инвентарем завершена",
            
            PharmacySavePoints.MEDICINE_START: "Начало операции с лекарством",
            PharmacySavePoints.MEDICINE_VALIDATED: "Лекарство валидировано",
            PharmacySavePoints.MEDICINE_SUPPLIER_CHECKED: "Поставщик проверен",
            PharmacySavePoints.MEDICINE_COMPLETED: "Операция с лекарством завершена",
            
            PharmacySavePoints.SUPPLIER_START: "Начало операции с поставщиком",
            PharmacySavePoints.SUPPLIER_VALIDATED: "Поставщик валидирован",
            PharmacySavePoints.SUPPLIER_COMPLETED: "Операция с поставщиком завершена"
        }
        return descriptions.get(savepoint_name, f"Точка сохранения: {savepoint_name}")


class DatabaseInitializerThread(QThread):
    finished = pyqtSignal(bool)
    progress = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.connection_params = {
            'dbname': 'pharmacydb',
            'user': 'pharmacy_admin',
            'password': 'pharmacyadmin',
            'host': 'localhost',
            'port': 5432,
            'options': '-c client_encoding=UTF8'
}

    def run(self):
        try:
            self.progress.emit("Подключение к базе данных...")
            conn = psycopg2.connect(**self.connection_params)
            conn.autocommit = True
            cursor = conn.cursor()

            # Создание ролей
            self.progress.emit("Создание ролей пользователей...")
            self._create_roles(cursor)

            # Создание таблиц
            self.progress.emit("Создание структуры таблиц...")
            self._create_tables(cursor)

            # Создание функций и триггеров
            self.progress.emit("Создание функций и триггеров...")
            self._create_functions_and_triggers(cursor)

            # Тестовые данные
            self.progress.emit("Добавление тестовых данных...")
            self._insert_test_data(cursor)

            cursor.close()
            conn.close()
            self.finished.emit(True)

        except Exception as e:
            print(f"Ошибка инициализации БД: {e}")
            self.finished.emit(False)

    def _create_roles(self, cursor):
        roles_sql = [
            """
            DO $$ BEGIN 
                IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'pharmacy_admin') THEN 
                    CREATE ROLE pharmacy_admin LOGIN PASSWORD 'pharmacyadmin'; 
                END IF; 
            END $$;
            """,
            "DO $$ BEGIN IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'pharmacy_user') THEN CREATE ROLE pharmacy_user LOGIN PASSWORD 'pharmacyuser'; END IF; END $$;",
            """
                            -- Создание роли для администратора базы данных
                            DO $$ BEGIN
                            IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'pharmacy_db_admin') THEN
                                CREATE ROLE pharmacy_db_admin WITH LOGIN PASSWORD 'pharmacyadmin2';
                            END IF;
                            END $$;
    
                            GRANT ALL PRIVILEGES ON DATABASE pharmacydb TO pharmacy_db_admin;
                            """,
            """
                            -- Создание роли для фармацевта
                            DO $$ BEGIN
                            IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'pharmacist') THEN
                                CREATE ROLE pharmacist WITH LOGIN PASSWORD 'pharmacist123';
                            END IF;
                            END $$;
    
                            GRANT SELECT, UPDATE, DELETE, INSERT ON ALL TABLES IN SCHEMA public TO pharmacist;
                            ALTER DEFAULT PRIVILEGES IN SCHEMA public GRANT SELECT, UPDATE, DELETE, INSERT ON TABLES TO pharmacist;
                            """,
            """
                            -- Создание роли кассира
                            DO $$ BEGIN
                            IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'cashier') THEN
                                CREATE ROLE cashier WITH LOGIN PASSWORD 'cashier123';
                            END IF;
                            END $$;
    
                            GRANT SELECT, INSERT ON Sales, SaleItems TO cashier;
                            GRANT SELECT ON Medicine, Inventory, Supplier TO cashier;
                            """,
        ]
        for sql in roles_sql:
            cursor.execute(sql)

    def _create_tables(self, cursor):
        tables_sql = [
            """
                -- Таблица пользователей системы
                CREATE TABLE IF NOT EXISTS public.dbusers (
                    ID_User INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    user_email VARCHAR(255) UNIQUE NOT NULL,
                    user_password VARCHAR(255) NOT NULL,
                    user_role VARCHAR(50) DEFAULT 'pharmacy_user',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );

                -- Таблица поставщиков
                CREATE TABLE IF NOT EXISTS public.Supplier (
                    ID_Supplier INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    SupplierName VARCHAR(255) NOT NULL,
                    ContactPerson VARCHAR(255),
                    Phone VARCHAR(50),
                    Email VARCHAR(255),
                    Address TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );

                -- Таблица лекарств
                CREATE TABLE IF NOT EXISTS public.Medicine (
                    ID_Medicine INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    MedicineName VARCHAR(255) NOT NULL,
                    ActiveSubstance VARCHAR(255),
                    Dosage VARCHAR(100),
                    Form VARCHAR(100), -- таблетки, сироп, мазь и т.д.
                    PrescriptionRequired BOOLEAN DEFAULT FALSE,
                    Price DECIMAL(10,2) NOT NULL,
                    ID_Supplier INT,
                    FOREIGN KEY (ID_Supplier) REFERENCES Supplier(ID_Supplier) ON DELETE SET NULL
                );

                -- Таблица инвентаря (остатки на складе)
                CREATE TABLE IF NOT EXISTS public.Inventory (
                    ID_Inventory INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    ID_Medicine INT NOT NULL,
                    BatchNumber VARCHAR(100),
                    Quantity INT NOT NULL DEFAULT 0,
                    ExpiryDate DATE,
                    PurchasePrice DECIMAL(10,2),
                    PurchaseDate DATE DEFAULT CURRENT_DATE,
                    FOREIGN KEY (ID_Medicine) REFERENCES Medicine(ID_Medicine) ON DELETE CASCADE
                );

                -- Таблица продаж
                CREATE TABLE IF NOT EXISTS public.Sales (
                    ID_Sale INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    SaleDate TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    CustomerName VARCHAR(255),
                    CustomerPhone VARCHAR(50),
                    TotalAmount DECIMAL(10,2) NOT NULL,
                    PaymentMethod VARCHAR(50) DEFAULT 'cash',
                    ID_User INT,
                    FOREIGN KEY (ID_User) REFERENCES dbusers(ID_User) ON DELETE SET NULL
                );

                -- Таблица позиций продаж
                CREATE TABLE IF NOT EXISTS public.SaleItems (
                    ID_SaleItem INT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                    ID_Sale INT NOT NULL,
                    ID_Medicine INT NOT NULL,
                    Quantity INT NOT NULL,
                    UnitPrice DECIMAL(10,2) NOT NULL,
                    TotalPrice DECIMAL(10,2) NOT NULL,
                    FOREIGN KEY (ID_Sale) REFERENCES Sales(ID_Sale) ON DELETE CASCADE,
                    FOREIGN KEY (ID_Medicine) REFERENCES Medicine(ID_Medicine) ON DELETE CASCADE
                );
                """
        ]
        for sql in tables_sql:
            cursor.execute(sql)

    def _create_functions_and_triggers(self, cursor):
        functions_sql = [
            """
            -- Функция для автоматического обновления остатков при продаже
            CREATE OR REPLACE FUNCTION update_inventory_on_sale()
            RETURNS TRIGGER AS $$
            BEGIN
                -- Уменьшаем количество в инвентаре при продаже
                UPDATE Inventory 
                SET Quantity = Quantity - NEW.Quantity
                WHERE ID_Medicine = NEW.ID_Medicine 
                AND Quantity >= NEW.Quantity
                AND ExpiryDate > CURRENT_DATE;
                
                -- Проверяем, что остаток достаточен
                IF NOT FOUND THEN
                    RAISE EXCEPTION 'Недостаточно товара на складе или товар просрочен';
                END IF;
                
                RETURN NEW;
            END;
            $$ LANGUAGE plpgsql;
            """,
            """
            -- Функция для проверки срока годности
            CREATE OR REPLACE FUNCTION check_expiry_date()
            RETURNS TRIGGER AS $$
            BEGIN
                -- Проверяем, что срок годности не истек
                IF NEW.ExpiryDate <= CURRENT_DATE THEN
                    RAISE EXCEPTION 'Товар с истекшим сроком годности не может быть добавлен';
                END IF;
                RETURN NEW;
            END;
            $$ LANGUAGE plpgsql;
            """,
            """
            -- Функция для автоматического расчета общей суммы продажи
            CREATE OR REPLACE FUNCTION calculate_sale_total()
            RETURNS TRIGGER AS $$
            DECLARE
                total_amount DECIMAL(10,2);
            BEGIN
                -- Рассчитываем общую сумму продажи
                SELECT COALESCE(SUM(TotalPrice), 0) INTO total_amount
                FROM SaleItems 
                WHERE ID_Sale = NEW.ID_Sale;
                
                -- Обновляем общую сумму в таблице Sales
                UPDATE Sales 
                SET TotalAmount = total_amount 
                WHERE ID_Sale = NEW.ID_Sale;
                
                RETURN NEW;
            END;
            $$ LANGUAGE plpgsql;
            """,
            """
            -- Представление для просмотра остатков с информацией о лекарствах
            CREATE OR REPLACE VIEW inventory_view AS
            SELECT 
                i.ID_Inventory,
                m.MedicineName,
                m.ActiveSubstance,
                m.Dosage,
                m.Form,
                i.BatchNumber,
                i.Quantity,
                i.ExpiryDate,
                i.PurchasePrice,
                s.SupplierName,
                CASE 
                    WHEN i.ExpiryDate <= CURRENT_DATE THEN 'Просрочен'
                    WHEN i.ExpiryDate <= CURRENT_DATE + INTERVAL '30 days' THEN 'Скоро истекает'
                    ELSE 'Годен'
                END as Status
            FROM Inventory i
            JOIN Medicine m ON i.ID_Medicine = m.ID_Medicine
            LEFT JOIN Supplier s ON m.ID_Supplier = s.ID_Supplier;
            """,
            """
            -- Триггер для обновления остатков при продаже
            CREATE TRIGGER trg_update_inventory_on_sale
            AFTER INSERT ON SaleItems
            FOR EACH ROW
            EXECUTE FUNCTION update_inventory_on_sale();
            """,
            """
            -- Триггер для проверки срока годности при добавлении в инвентарь
            CREATE TRIGGER trg_check_expiry_date
            BEFORE INSERT OR UPDATE ON Inventory
            FOR EACH ROW
            EXECUTE FUNCTION check_expiry_date();
            """,
            """
            -- Триггер для автоматического расчета суммы продажи
            CREATE TRIGGER trg_calculate_sale_total
            AFTER INSERT OR UPDATE OR DELETE ON SaleItems
            FOR EACH ROW
            EXECUTE FUNCTION calculate_sale_total();
            """
        ]
        for sql in functions_sql:
            cursor.execute(sql)

    def _insert_test_data(self, cursor):
        test_data_sql = [
            """
            INSERT INTO dbusers (user_email, user_password, user_role)
            VALUES 
                ('admin@pharmacy.com', 'admin123', 'pharmacy_admin'),
                ('pharmacist@pharmacy.com', 'pharmacist123', 'pharmacist'),
                ('cashier@pharmacy.com', 'cashier123', 'cashier')
            ON CONFLICT (user_email) DO NOTHING;
            """,
            """
            INSERT INTO Supplier (SupplierName, ContactPerson, Phone, Email, Address)
            VALUES 
                ('Фармацевтическая компания "Здоровье"', 'Иванов И.И.', '+7(495)123-45-67', 'info@zdorovie.ru', 'Москва, ул. Лекарственная, 1'),
                ('Медицинские препараты ООО', 'Петрова А.С.', '+7(495)234-56-78', 'sales@medprep.ru', 'СПб, пр. Медицинский, 15'),
                ('Аптечная сеть "Витамин"', 'Сидоров П.В.', '+7(495)345-67-89', 'supply@vitamin.ru', 'Казань, ул. Аптечная, 25')
            ON CONFLICT DO NOTHING;
            """,
            """
            INSERT INTO Medicine (MedicineName, ActiveSubstance, Dosage, Form, PrescriptionRequired, Price, ID_Supplier)
            VALUES 
                ('Аспирин', 'Ацетилсалициловая кислота', '500мг', 'таблетки', false, 150.00, 1),
                ('Парацетамол', 'Парацетамол', '500мг', 'таблетки', false, 120.00, 1),
                ('Амоксициллин', 'Амоксициллин', '250мг', 'капсулы', true, 350.00, 2),
                ('Витамин С', 'Аскорбиновая кислота', '100мг', 'таблетки', false, 80.00, 3),
                ('Нурофен', 'Ибупрофен', '200мг', 'таблетки', false, 280.00, 1)
            ON CONFLICT DO NOTHING;
            """,
            """
            INSERT INTO Inventory (ID_Medicine, BatchNumber, Quantity, ExpiryDate, PurchasePrice, PurchaseDate)
            VALUES 
                (1, 'ASP001', 100, '2025-12-31', 120.00, '2024-01-15'),
                (2, 'PAR002', 150, '2025-10-15', 95.00, '2024-01-20'),
                (3, 'AMX003', 50, '2025-08-30', 280.00, '2024-02-01'),
                (4, 'VIT004', 200, '2025-06-20', 65.00, '2024-01-25'),
                (5, 'NUR005', 75, '2025-11-10', 220.00, '2024-02-05')
            ON CONFLICT DO NOTHING;
            """
        ]
        for sql in test_data_sql:
            cursor.execute(sql)


class LoginWindow(QDialog):
    def __init__(self):
        super().__init__()
        self.registration_window = None
        self.conn = None
        self.cursor = None
        self.user_role = None
        self.init_db_thread = None

        loadUi(resource_path("enter.ui"), self)
        self.setFixedSize(self.size())

        self.pushButton_2 = self.findChild(QPushButton, "pushButton_2")
        self.pushButton_3 = self.findChild(QPushButton, "pushButton_3")
        self.lineEdit = self.findChild(QLineEdit, "lineEdit")
        self.lineEdit_2 = self.findChild(QLineEdit, "lineEdit_2")

        self.lineEdit_2.setEchoMode(QLineEdit.EchoMode.Password)

        self.pushButton_2.clicked.connect(lambda: self.handle_login())
        self.pushButton_3.clicked.connect(lambda: self.handle_register())

    @with_sound("click")
    def handle_login(self):
        """Обработчик для кнопки входа"""
        self.check_database_connection()

    @with_sound("click")
    def handle_register(self):
        """Обработчик для кнопки регистрации"""
        self.open_register_window()

    def check_database_connection(self):
        try:
            test_conn = psycopg2.connect(
                dbname='pharmacydb',
                user='pharmacy_admin',
                password='pharmacyadmin',
                host='localhost',
                connect_timeout=3
            )
            test_conn.close()
            self.connect_to_db()
        except Exception as e:
            QMessageBox.information(self, "Информация",
                                    "База данных не доступна. Попытка создания...")
            self.create_database_if_not_exists()

    def create_database_if_not_exists(self):
        try:
            # Подключаемся суперпользователем (postgres)
            admin_conn = psycopg2.connect(
                dbname='postgres',
                user='postgres',
                password='postgres',
                host='localhost'
            )
            admin_conn.autocommit = True
            cursor = admin_conn.cursor()

            # Проверяем наличие базы
            cursor.execute("SELECT 1 FROM pg_database WHERE datname='pharmacydb'")
            if not cursor.fetchone():
                cursor.execute("CREATE DATABASE pharmacydb")
                QMessageBox.information(self, "Успех", "База данных создана")

            # Создаём роль pharmacy_admin, если её нет
            cursor.execute("""
                DO $$
                BEGIN
                    IF NOT EXISTS (SELECT FROM pg_roles WHERE rolname = 'pharmacy_admin') THEN
                        CREATE ROLE pharmacy_admin LOGIN PASSWORD 'pharmacyadmin';
                    END IF;
                END
                $$;
            """)

            # Даём права на базу
            cursor.execute("GRANT ALL PRIVILEGES ON DATABASE pharmacydb TO pharmacy_admin;")
            cursor.execute("GRANT ALL PRIVILEGES ON DATABASE pharmacydb TO postgres;")

            cursor.close()
            admin_conn.close()

            # Запускаем инициализацию (создание таблиц, ролей, тестовых данных)
            self.init_db_thread = DatabaseInitializerThread()
            self.init_db_thread.progress.connect(self._show_progress_message)
            self.init_db_thread.finished.connect(self.on_database_initialized)
            self.init_db_thread.start()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось создать БД: {str(e)}")

    def _show_progress_message(self, message):
        print(f"Прогресс: {message}")  # Можно заменить на вывод в статус-бар

    def on_database_initialized(self, success):
        if success:
            self.connect_to_db()
        else:
            QMessageBox.critical(self, "Ошибка", "Не удалось инициализировать БД")

    def connect_to_db(self):
        user_input = self.lineEdit.text()
        password_input = self.lineEdit_2.text()

        try:
            if user_input == "pharmacy_admin" and password_input == "pharmacyadmin":
                self.user_role = "pharmacy_admin"
                self.conn = psycopg2.connect(
                    dbname='pharmacydb',
                    user='pharmacy_admin',
                    password='pharmacyadmin',
                    host='localhost',
                    port=5432
                )
                self.conn.autocommit = True
                self.cursor = self.conn.cursor()
                self.open_user_window()
                return

            with psycopg2.connect(
                    dbname='pharmacydb',
                    user='pharmacy_admin',
                    password='pharmacyadmin',
                    host='localhost',
                    port=5432
            ) as admin_conn:
                admin_cursor = admin_conn.cursor()
                admin_cursor.execute(
                    "SELECT user_password FROM dbusers WHERE user_email = %s",
                    (user_input,)
                )
                result = admin_cursor.fetchone()

                if result is not None:
                    stored_password = str(result[0])
                    if password_input == stored_password:
                        self.user_role = "pharmacy_user"
                        self.conn = psycopg2.connect(
                            dbname='pharmacydb',
                            user='pharmacy_user',
                            password='pharmacyuser',
                            host='localhost',
                            port=5432
                        )
                        self.conn.autocommit = True
                        self.cursor = self.conn.cursor()
                        self.open_user_window()
                    else:
                        QMessageBox.warning(self, "Ошибка", "Неверный пароль.")
                else:
                    QMessageBox.warning(self, "Ошибка", "Пользователь не найден.")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def open_register_window(self):
        self.user_window = RegisterWindow(self.conn, self.cursor)
        self.user_window.show()

    def open_user_window(self):
        self.user_window = UserWindow(self.conn, self.cursor, self.user_role)
        self.user_window.show()
        self.hide()



class RegisterWindow(QDialog):
    def __init__(self, conn, cursor):
        super().__init__()
        loadUi(resource_path("registration.ui"), self)

            # виджеты
        self.lineEdit = self.findChild(QLineEdit, "lineEdit")  # Почта
        self.lineEdit_2 = self.findChild(QLineEdit, "lineEdit_2")  # Пароль
        self.lineEdit_3 = self.findChild(QLineEdit, "lineEdit_3")  # Повтор пароля
        self.pushButton_2 = self.findChild(QPushButton, "pushButton_2")  # Кнопка регистрации

        self.lineEdit_2.setEchoMode(QLineEdit.EchoMode.Password)
        self.lineEdit_3.setEchoMode(QLineEdit.EchoMode.Password)

        self.setFixedSize(self.size())

        # Подключение сигналов
        self.pushButton_2.clicked.connect(self.register_with_sound)

    @with_sound("click")
    def register_with_sound(self):
        """Обработчик регистрации с звуком"""
        self.register_user()

    def register_user(self):



class SaleData:
    def __init__(self, sale_id, sale_date, customer_name, customer_phone, total_amount, payment_method, items):
        self.sale_id = sale_id
        self.sale_date = sale_date
        self.customer_name = customer_name
        self.customer_phone = customer_phone
        self.total_amount = total_amount
        self.payment_method = payment_method
        self.items = items  # Список позиций продажи

class MedicineData:
    def __init__(self, medicine_id, medicine_name, active_substance, dosage, form, prescription_required, price, supplier_name):
        self.medicine_id = medicine_id
        self.medicine_name = medicine_name
        self.active_substance = active_substance
        self.dosage = dosage
        self.form = form
        self.prescription_required = prescription_required
        self.price = price
        self.supplier_name = supplier_name


class UserWindow(QDialog):
    def __init__(self, conn, cursor, user_role):
        super().__init__()
        self.conn = conn
        self.cursor = cursor
        self.user_role = user_role
        self.transaction_manager = PharmacyTransactionManager(conn, cursor)

        loadUi(resource_path("pharmacy_main.ui"), self)
        self.setFixedSize(self.size())

        # Подключение виджетов из нового UI для аптеки
        self.medicine_combo = self.findChild(QComboBox, "comboBox_medicine")  # Лекарство
        self.date_edit = self.findChild(QDateEdit, "dateEdit")  # Дата
        self.search_button = self.findChild(QPushButton, "pushButton_search")  # Поиск
        self.add_sale_button = self.findChild(QPushButton, "pushButton_add_sale")  # Добавить продажу
        
        # Таблицы из TabWidget
        self.tab_widget = self.findChild(QTabWidget, "tabWidget")
        self.table_sales = self.findChild(QTableView, "tableView_sales")  # Таблица продаж
        self.table_inventory = self.findChild(QTableView, "tableView_inventory")  # Таблица инвентаря
        self.table_expiring = self.findChild(QTableView, "tableView_expiring")  # Таблица товаров с истекающим сроком
        self.table_medicines = self.findChild(QTableView, "tableView_medicines")  # Таблица лекарств

        self.supplier_combo = self.findChild(QComboBox, "comboBox_supplier")  # Поставщик
        self.status_label = self.findChild(QLabel, "statusLabel")  # Статус бар

        self.checkbox_medicine = self.findChild(QCheckBox, "checkBox_medicine")
        self.checkbox_supplier = self.findChild(QCheckBox, "checkBox_supplier")
        self.checkbox_date = self.findChild(QCheckBox, "checkBox_date")

        # Инициализация чекбоксов
        self.checkbox_medicine.setChecked(False)
        self.checkbox_supplier.setChecked(False)
        self.checkbox_date.setChecked(False)

        # Скрываем элементы фильтров по умолчанию
        self.medicine_combo.setVisible(False)
        self.supplier_combo.setVisible(False)
        self.date_edit.setVisible(False)

        # Подключение сигналов чекбоксов
        self.checkbox_medicine.stateChanged.connect(self.handle_checkbox_medicine)
        self.checkbox_supplier.stateChanged.connect(self.handle_checkbox_supplier)
        self.checkbox_date.stateChanged.connect(self.handle_checkbox_date)
        
        # Проверка роли пользователя
        if self.user_role not in ["pharmacy_admin", "pharmacist", "cashier"]:
            self.add_sale_button.hide()

        # Подключение кнопок
        self.search_button.clicked.connect(lambda: self.handle_search_button_click())
        self.add_sale_button.clicked.connect(lambda: self.handle_add_sale_button_click())
        
        # Подключение двойных кликов по таблицам
        self.table_sales.doubleClicked.connect(lambda: self.handle_sales_double_click())
        self.table_inventory.doubleClicked.connect(lambda: self.handle_inventory_double_click())
        self.table_expiring.doubleClicked.connect(lambda: self.handle_expiring_double_click())
        self.table_medicines.doubleClicked.connect(lambda: self.handle_medicines_double_click())

        # Загрузка данных
        self.load_medicines()
        self.load_suppliers()
        
        # Первоначальная загрузка данных
        QTimer.singleShot(0, self.load_initial_data)
        
        # Обновление статуса
        self.update_status("Система готова к работе")

    @with_sound("click")
    def handle_search_button_click(self):
        """Обработчик клика по кнопке поиска"""
        self.load_sales_data()

    @with_sound("click")
    def handle_add_sale_button_click(self):
        """Обработчик клика по кнопке добавления продажи"""
        self.add_sale()

    @with_sound("click")
    def handle_sales_double_click(self, index):
        """Обработчик двойного клика по таблице продаж"""
        self.show_sale_details(index, self.table_sales)

    @with_sound("click")
    def handle_inventory_double_click(self, index):
        """Обработчик двойного клика по таблице инвентаря"""
        self.show_inventory_details(index, self.table_inventory)

    @with_sound("click")
    def handle_expiring_double_click(self, index):
        """Обработчик двойного клика по таблице товаров с истекающим сроком"""
        self.show_expiring_details(index, self.table_expiring)

    @with_sound("click")
    def handle_medicines_double_click(self, index):
        """Обработчик двойного клика по таблице лекарств"""
        self.show_medicine_details(index, self.table_medicines)

    @with_sound("click")
    def handle_checkbox_medicine(self, state):
        """Обработчик для чекбокса лекарства"""
        self.toggle_medicine_visibility(state)

    @with_sound("click")
    def handle_checkbox_supplier(self, state):
        """Обработчик для чекбокса поставщика"""
        self.toggle_supplier_visibility(state)

    @with_sound("click")
    def handle_checkbox_date(self, state):
        """Обработчик для чекбокса даты"""
        self.toggle_date_visibility(state)

    @with_sound("click")
    def handle_result_button_click(self):
        """Обработчик клика по кнопке результатов"""
        # Сначала загружаем данные
        self.load_data()

        # Проверяем, есть ли данные для сохранения
        model = self.table_view.model()
        if not model or model.rowCount() == 0:
            QMessageBox.information(self, "Информация", "Нет данных для сохранения")
            return

        # Запрашиваем пользователя о сохранении
        reply = QMessageBox.question(
            self,
            "Сохранение результатов",
            "Хотите сохранить результаты в файл?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes
        )

        if reply == QMessageBox.StandardButton.Yes:
            # Получаем данные из модели
            data = []
            for row in range(model.rowCount()):
                row_data = []
                for column in range(model.columnCount()):
                    index = model.index(row, column)
                    row_data.append(model.data(index))
                data.append(row_data)

            # Вызываем экспорт
            self.export_to_docx(data)

    def check_connection(self):
        try:
            self.cursor.execute("SELECT 1")
            return True
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Потеряно соединение с базой данных:\n{e}")
            return False

    def load_play_types(self):
        try:
            self.cursor.execute("SELECT DISTINCT SportName FROM SportType;")
            plays = self.cursor.fetchall()
            self.play_combo.clear()
            self.play_combo.addItem("Все")
            for play in plays:
                self.play_combo.addItem(play[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки списка игр:\n{e}")

    def load_tournaments_for_filter(self):
        try:
            self.cursor.execute("SELECT DISTINCT TournamentName FROM Tournament;")
            tournaments = self.cursor.fetchall()
            self.tournament_combo_5.clear()
            self.tournament_combo_5.addItem("Все")
            for tournament in tournaments:
                self.tournament_combo_5.addItem(tournament[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки списка турниров:\n{e}")

    def toggle_medicine_visibility(self, state):
        """Переключение видимости фильтра лекарств"""
        visible = state == 2  # Qt.Checked.value
        self.medicine_combo.setVisible(visible)

    def toggle_supplier_visibility(self, state):
        """Переключение видимости фильтра поставщиков"""
        visible = state == 2
        self.supplier_combo.setVisible(visible)

    def toggle_date_visibility(self, state):
        """Переключение видимости фильтра даты"""
        visible = state == 2
        self.date_edit.setVisible(visible)

    def update_status(self, message):
        """Обновление статус бара"""
        if hasattr(self, 'status_label') and self.status_label:
            self.status_label.setText(message)
        print(f"Статус: {message}")

    def load_medicines(self):
        """Загрузка списка лекарств"""
        try:
            self.cursor.execute("SELECT MedicineName FROM Medicine ORDER BY MedicineName;")
            medicines = self.cursor.fetchall()
            self.medicine_combo.clear()
            self.medicine_combo.addItem("Все лекарства")
            for medicine in medicines:
                self.medicine_combo.addItem(medicine[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки лекарств:\n{e}")

    def load_suppliers(self):
        """Загрузка списка поставщиков"""
        try:
            self.cursor.execute("SELECT SupplierName FROM Supplier ORDER BY SupplierName;")
            suppliers = self.cursor.fetchall()
            self.supplier_combo.clear()
            self.supplier_combo.addItem("Все поставщики")
            for supplier in suppliers:
                self.supplier_combo.addItem(supplier[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки поставщиков:\n{e}")

    def load_initial_data(self):
        """Первоначальная загрузка всех данных"""
        try:
            self.update_status("Загрузка данных...")
            self.load_sales_data()
            self.load_inventory_data()
            self.load_expiring_data()
            self.load_medicines_data()
            self.update_status("Данные загружены успешно")
        except Exception as e:
            self.update_status(f"Ошибка загрузки: {str(e)}")
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки данных:\n{e}")

    def load_sales_data(self):
        """Загрузка данных о продажах"""
        try:
            # Формируем условия фильтрации
            conditions = []
            params = []
            
            if self.checkbox_medicine.isChecked() and self.medicine_combo.currentText() != "Все лекарства":
                conditions.append("m.MedicineName = %s")
                params.append(self.medicine_combo.currentText())
                
            if self.checkbox_date.isChecked() and self.date_edit.date().isValid():
                selected_date = self.date_edit.date().toString("yyyy-MM-dd")
                conditions.append("DATE(s.SaleDate) = %s")
                params.append(selected_date)
            
            # Базовый запрос
            base_query = """
                SELECT 
                    s.ID_Sale,
                    s.SaleDate,
                    s.CustomerName,
                    s.CustomerPhone,
                    s.TotalAmount,
                    s.PaymentMethod,
                    COUNT(si.ID_SaleItem) as ItemsCount
                FROM Sales s
                LEFT JOIN SaleItems si ON s.ID_Sale = si.ID_Sale
                LEFT JOIN Medicine m ON si.ID_Medicine = m.ID_Medicine
                WHERE 1=1
            """
            
            if conditions:
                full_query = base_query + " AND " + " AND ".join(conditions)
            else:
                full_query = base_query
                
            full_query += " GROUP BY s.ID_Sale, s.SaleDate, s.CustomerName, s.CustomerPhone, s.TotalAmount, s.PaymentMethod ORDER BY s.SaleDate DESC LIMIT 100"
            
            self.cursor.execute(full_query, params)
            sales_data = self.cursor.fetchall()
            
            # Форматируем данные для таблицы
            formatted_data = []
            for row in sales_data:
                formatted_data.append((
                    str(row[0]),  # ID продажи
                    row[1].strftime("%Y-%m-%d %H:%M") if hasattr(row[1], 'strftime') else str(row[1]),  # Дата
                    row[2] or "Не указано",  # Покупатель
                    row[3] or "Не указано",  # Телефон
                    f"{row[4]:.2f} руб.",  # Сумма
                    row[5],  # Способ оплаты
                    str(row[6])  # Количество позиций
                ))
            
            headers = ["ID", "Дата", "Покупатель", "Телефон", "Сумма", "Оплата", "Позиций"]
            self.update_table(self.table_sales, formatted_data, headers)
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки продаж:\n{e}")

    def load_inventory_data(self):
        """Загрузка данных об инвентаре"""
        try:
            query = """
                SELECT 
                    m.MedicineName,
                    m.ActiveSubstance,
                    m.Dosage,
                    m.Form,
                    i.BatchNumber,
                    i.Quantity,
                    i.ExpiryDate,
                    i.PurchasePrice,
                    s.SupplierName,
                    CASE 
                        WHEN i.ExpiryDate <= CURRENT_DATE THEN 'Просрочен'
                        WHEN i.ExpiryDate <= CURRENT_DATE + INTERVAL '30 days' THEN 'Скоро истекает'
                        ELSE 'Годен'
                    END as Status
                FROM Inventory i
                JOIN Medicine m ON i.ID_Medicine = m.ID_Medicine
                LEFT JOIN Supplier s ON m.ID_Supplier = s.ID_Supplier
                ORDER BY i.ExpiryDate ASC
            """
            
            self.cursor.execute(query)
            inventory_data = self.cursor.fetchall()
            
            # Форматируем данные для таблицы
            formatted_data = []
            for row in inventory_data:
                formatted_data.append((
                    row[0],  # Название лекарства
                    row[1] or "Не указано",  # Действующее вещество
                    row[2] or "Не указано",  # Дозировка
                    row[3] or "Не указано",  # Форма
                    row[4] or "Не указано",  # Номер партии
                    str(row[5]),  # Количество
                    row[6].strftime("%Y-%m-%d") if hasattr(row[6], 'strftime') else str(row[6]),  # Срок годности
                    f"{row[7]:.2f} руб." if row[7] else "Не указано",  # Цена закупки
                    row[8] or "Не указано",  # Поставщик
                    row[9]  # Статус
                ))
            
            headers = ["Лекарство", "Действующее вещество", "Дозировка", "Форма", "Партия", "Количество", "Срок годности", "Цена закупки", "Поставщик", "Статус"]
            self.update_table(self.table_inventory, formatted_data, headers)
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки инвентаря:\n{e}")

    def load_expiring_data(self):
        """Загрузка товаров с истекающим сроком годности"""
        try:
            query = """
                SELECT 
                    m.MedicineName,
                    m.ActiveSubstance,
                    m.Dosage,
                    m.Form,
                    i.BatchNumber,
                    i.Quantity,
                    i.ExpiryDate,
                    s.SupplierName,
                    CASE 
                        WHEN i.ExpiryDate <= CURRENT_DATE THEN 'Просрочен'
                        WHEN i.ExpiryDate <= CURRENT_DATE + INTERVAL '30 days' THEN 'Скоро истекает'
                        ELSE 'Годен'
                    END as Status
                FROM Inventory i
                JOIN Medicine m ON i.ID_Medicine = m.ID_Medicine
                LEFT JOIN Supplier s ON m.ID_Supplier = s.ID_Supplier
                WHERE i.ExpiryDate <= CURRENT_DATE + INTERVAL '30 days'
                ORDER BY i.ExpiryDate ASC
            """
            
            self.cursor.execute(query)
            expiring_data = self.cursor.fetchall()
            
            # Форматируем данные для таблицы
            formatted_data = []
            for row in expiring_data:
                formatted_data.append((
                    row[0],  # Название лекарства
                    row[1] or "Не указано",  # Действующее вещество
                    row[2] or "Не указано",  # Дозировка
                    row[3] or "Не указано",  # Форма
                    row[4] or "Не указано",  # Номер партии
                    str(row[5]),  # Количество
                    row[6].strftime("%Y-%m-%d") if hasattr(row[6], 'strftime') else str(row[6]),  # Срок годности
                    row[7] or "Не указано",  # Поставщик
                    row[8]  # Статус
                ))
            
            headers = ["Лекарство", "Действующее вещество", "Дозировка", "Форма", "Партия", "Количество", "Срок годности", "Поставщик", "Статус"]
            self.update_table(self.table_expiring, formatted_data, headers)
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки товаров с истекающим сроком:\n{e}")

    def load_medicines_data(self):
        """Загрузка данных о лекарствах"""
        try:
            query = """
                SELECT 
                    m.MedicineName,
                    m.ActiveSubstance,
                    m.Dosage,
                    m.Form,
                    m.PrescriptionRequired,
                    m.Price,
                    s.SupplierName,
                    COALESCE(SUM(i.Quantity), 0) as TotalQuantity
                FROM Medicine m
                LEFT JOIN Supplier s ON m.ID_Supplier = s.ID_Supplier
                LEFT JOIN Inventory i ON m.ID_Medicine = i.ID_Medicine AND i.ExpiryDate > CURRENT_DATE
                GROUP BY m.ID_Medicine, m.MedicineName, m.ActiveSubstance, m.Dosage, m.Form, m.PrescriptionRequired, m.Price, s.SupplierName
                ORDER BY m.MedicineName
            """
            
            self.cursor.execute(query)
            medicines_data = self.cursor.fetchall()
            
            # Форматируем данные для таблицы
            formatted_data = []
            for row in medicines_data:
                formatted_data.append((
                    row[0],  # Название лекарства
                    row[1] or "Не указано",  # Действующее вещество
                    row[2] or "Не указано",  # Дозировка
                    row[3] or "Не указано",  # Форма
                    "Да" if row[4] else "Нет",  # Требуется рецепт
                    f"{row[5]:.2f} руб.",  # Цена
                    row[6] or "Не указано",  # Поставщик
                    str(row[7])  # Общее количество на складе
                ))
            
            headers = ["Лекарство", "Действующее вещество", "Дозировка", "Форма", "Рецепт", "Цена", "Поставщик", "На складе"]
            self.update_table(self.table_medicines, formatted_data, headers)
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки лекарств:\n{e}")

    def update_table(self, table_view, data, headers=None):
        try:
            model = QStandardItemModel()

            if headers:
                model.setHorizontalHeaderLabels(headers)

            # Удаление дубликатов
            unique_data = []
            seen = set()
            for row in data:
                row_tuple = tuple(row)
                if row_tuple not in seen:
                    seen.add(row_tuple)
                    unique_data.append(row)

            model._table_data = unique_data

            for row in unique_data:
                items = [QStandardItem(str(item) if item is not None else "") for item in row]
                model.appendRow(items)

            table_view.setModel(model)
            table_view.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
            self.connect_table_signals()

        except Exception as e:
            print(f"Ошибка обновления таблицы: {e}")
            QMessageBox.critical(self, "Ошибка", f"Ошибка обновления таблицы:\n{e}")

    def load_recent_and_upcoming_matches(self):
        try:
            if not hasattr(self, 'check_connection') or not self.check_connection():
                return

            selected_sport = self.play_combo.currentText()
            selected_tournament = self.tournament_combo_5.currentText()

            # Условия фильтрации
            sport_condition = "AND st.SportName = %s" if selected_sport != "Все" else ""
            tournament_condition = "AND tr.TournamentName = %s" if selected_tournament != "Все" else ""

            # Параметры запроса
            params = []
            if selected_sport != "Все":
                params.append(selected_sport)
            if selected_tournament != "Все":
                params.append(selected_tournament)

            # Запрос для завершенных матчей
            # Запрос для завершенных матчей
            query_recent = f"""
            SELECT 
                m.MatchDateTime,
                t1.TeamName AS HomeTeam,
                t2.TeamName AS AwayTeam,
                r1.Score AS HomeScore,
                r2.Score AS AwayScore,
                st.SportName,
                COALESCE(tr.TournamentName, 'Не указан') AS TournamentName
            FROM Match m
            JOIN Team t1 ON m.ID_Team1 = t1.ID_Team
            JOIN Team t2 ON m.ID_Team2 = t2.ID_Team
            JOIN SportType st ON m.ID_SportType = st.ID_SportType
            LEFT JOIN (
                SELECT r.ID_Match, r.ID_Team, r.Score
                FROM Result r
                JOIN (
                    SELECT ID_Match, ID_Team, MAX(ID_Result) as MaxID
                    FROM Result
                    GROUP BY ID_Match, ID_Team
                ) latest ON r.ID_Match = latest.ID_Match 
                          AND r.ID_Team = latest.ID_Team 
                          AND r.ID_Result = latest.MaxID
            ) r1 ON r1.ID_Match = m.ID_Match AND r1.ID_Team = t1.ID_Team
            LEFT JOIN (
                SELECT r.ID_Match, r.ID_Team, r.Score
                FROM Result r
                JOIN (
                    SELECT ID_Match, ID_Team, MAX(ID_Result) as MaxID
                    FROM Result
                    GROUP BY ID_Match, ID_Team
                ) latest ON r.ID_Match = latest.ID_Match 
                          AND r.ID_Team = latest.ID_Team 
                          AND r.ID_Result = latest.MaxID
            ) r2 ON r2.ID_Match = m.ID_Match AND r2.ID_Team = t2.ID_Team
            LEFT JOIN Tournament tr ON m.ID_Tournament = tr.ID_Tournament
            WHERE m.MatchDateTime < CURRENT_TIMESTAMP
            {sport_condition}
            {tournament_condition}
            ORDER BY m.MatchDateTime DESC
            LIMIT 10
            """

            # Упрощенный запрос для будущих матчей
            query_upcoming = f"""
                    SELECT 
                        m.MatchDateTime,
                        t1.TeamName AS HomeTeam,
                        t2.TeamName AS AwayTeam,
                        NULL AS HomeScore,
                        NULL AS AwayScore,
                        st.SportName,
                        COALESCE(tr.TournamentName, 'Не указан') AS TournamentName
                    FROM Match m
                    JOIN Team t1 ON m.ID_Team1 = t1.ID_Team
                    JOIN Team t2 ON m.ID_Team2 = t2.ID_Team
                    JOIN SportType st ON m.ID_SportType = st.ID_SportType
                    LEFT JOIN Tournament tr ON m.ID_Tournament = tr.ID_Tournament
                    WHERE m.MatchDateTime > CURRENT_TIMESTAMP
                    {sport_condition}
                    {tournament_condition}
                    ORDER BY m.MatchDateTime ASC
                    LIMIT 10
                    """

            # Выполняем запросы
            self.cursor.execute(query_recent, params)
            recent_matches = []
            for row in self.cursor.fetchall():
                recent_matches.append((
                    row[0].strftime("%Y-%m-%d %H:%M"),
                    row[1], row[2],
                    str(row[3]) if row[3] is not None else "-",
                    str(row[4]) if row[4] is not None else "-",
                    row[5], row[6]
                ))

            self.cursor.execute(query_upcoming, params)
            upcoming_matches = []
            for row in self.cursor.fetchall():
                upcoming_matches.append((
                    row[0].strftime("%Y-%m-%d %H:%M"),
                    row[1], row[2],
                    "-", "-",  # Для предстоящих матчей
                    row[5], row[6]
                ))

            headers = ["Дата", "Команда 1", "Команда 2", "Счёт 1", "Счёт 2", "Вид спорта", "Турнир"]
            self.update_table(self.table_last_games, recent_matches, headers)
            self.update_table(self.table_future_games, upcoming_matches, headers)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки матчей:\n{str(e)}")
            print(f"Ошибка: {e}")
            import traceback
            traceback.print_exc()

    def load_data(self):
        try:
            if not hasattr(self, 'check_connection') or not self.check_connection():
                return

            # Базовый запрос с исправленными подзапросами
            base_query = """
            SELECT 
                m.MatchDateTime, 
                t1.TeamName AS Team1, 
                t2.TeamName AS Team2,
                r1.Score AS Score1,
                r2.Score AS Score2,
                st.SportName,
                COALESCE(tr.TournamentName, 'Не указан') AS TournamentName
            FROM Match m
            JOIN Team t1 ON m.ID_Team1 = t1.ID_Team
            JOIN Team t2 ON m.ID_Team2 = t2.ID_Team
            JOIN SportType st ON m.ID_SportType = st.ID_SportType
            LEFT JOIN Result r1 ON r1.ID_Match = m.ID_Match AND r1.ID_Team = t1.ID_Team
            LEFT JOIN Result r2 ON r2.ID_Match = m.ID_Match AND r2.ID_Team = t2.ID_Team
            LEFT JOIN Tournament tr ON m.ID_Tournament = tr.ID_Tournament
            WHERE 1=1
            """

            conditions = []
            params = []

            # Условия фильтрации
            if self.checkbox_sport_type.isChecked():
                conditions.append("st.SportName = %s")
                params.append(self.sport_combo.currentText())

            if self.checkbox_tournament.isChecked() and self.tournament_combo.currentText() != "Все турниры":
                conditions.append("tr.TournamentName = %s")
                params.append(self.tournament_combo.currentText())

            if self.checkbox_team.isChecked() and self.team_combo.currentText() != "Все команды":
                conditions.append("(t1.TeamName = %s OR t2.TeamName = %s)")
                params.extend([self.team_combo.currentText(), self.team_combo.currentText()])

            if self.checkbox_date.isChecked() and self.date_edit.date().isValid():
                selected_date = self.date_edit.date().toString("yyyy-MM-dd")
                conditions.append("DATE(m.MatchDateTime) = %s")
                params.append(selected_date)

            if self.play_combo.currentText() != "Все":
                conditions.append("st.SportName = %s")
                params.append(self.play_combo.currentText())

            if self.tournament_combo_5.currentText() != "Все":
                conditions.append("tr.TournamentName = %s")
                params.append(self.tournament_combo_5.currentText())

            # Формируем полный запрос
            if conditions:
                full_query = base_query + " AND " + " AND ".join(conditions)
            else:
                full_query = base_query

            full_query += " ORDER BY m.MatchDateTime DESC"

            # Выполняем запрос
            self.cursor.execute(full_query, params)
            all_data = self.cursor.fetchall()

            # Форматируем данные для таблицы
            matches = []
            for row in all_data:
                matches.append((
                    row[0].strftime("%Y-%m-%d %H:%M") if hasattr(row[0], 'strftime') else str(row[0]),
                    row[1], row[2],
                    str(row[3]) if row[3] is not None else "-",
                    str(row[4]) if row[4] is not None else "-",
                    row[5], row[6]
                ))

            headers = ["Дата", "Команда 1", "Команда 2", "Счёт 1", "Счёт 2", "Вид спорта", "Турнир"]
            self.update_table(self.table_view, matches, headers)

        except Exception as e:
            print(f"Ошибка загрузки данных: {str(e)}")
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки данных:\n{str(e)}")

    def show_team_players(self, team_name):
        try:
            self.cursor.execute("""
                SELECT p.FirstName, p.LastName, p.Position, p.DateOfBirth
                FROM Player p
                JOIN Team t ON p.ID_Team = t.ID_Team
                WHERE t.TeamName = %s;
            """, (team_name,))
            players = self.cursor.fetchall()

            if players:
                model = QStandardItemModel(self)
                model.setColumnCount(4)
                model.setHorizontalHeaderLabels(["Имя", "Фамилия", "Позиция", "Дата рождения"])

                for row in players:
                    items = [QStandardItem(str(cell)) for cell in row]
                    model.appendRow(items)

                self.table_view.setModel(model)
            else:
                QMessageBox.information(self, "Информация", "У этой команды нет игроков.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки игроков:\n{e}")

    def export_to_docx(self, data):
        try:
            doc = Document()
            doc.add_heading('Результаты матчей', 0)

            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Команда 1'
            hdr_cells[2].text = 'Счёт 1'
            hdr_cells[3].text = 'Счёт 2'
            hdr_cells[4].text = 'Команда 2'

            for row in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row[0])
                row_cells[1].text = str(row[1])
                row_cells[2].text = str(row[3]) if row[3] is not None else "-"
                row_cells[3].text = str(row[4]) if row[4] is not None else "-"
                row_cells[4].text = str(row[2])

            # Сохраняем через диалог
            filename, _ = QFileDialog.getSaveFileName(self, "Сохранить как", "", "Word Documents (*.docx)")
            if filename:
                if not filename.endswith(".docx"):
                    filename += ".docx"
                doc.save(filename)
                QMessageBox.information(self, "Успешно", "Документ сохранён.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить документ:\n{e}")

    def connect_table_signals(self):
        """Подключает сигналы двойного клика для всех таблиц"""
        # Отключаем все предыдущие соединения (если есть)
        try:
            self.table_view.doubleClicked.disconnect()
        except TypeError:
            pass

        try:
            self.table_last_games.doubleClicked.disconnect()
        except TypeError:
            pass

        try:
            self.table_future_games.doubleClicked.disconnect()
        except TypeError:
            pass

        # Подключаем новые обработчики с правильной сигнатурой для PyQt6
        self.table_view.doubleClicked.connect(
            lambda index: self.show_match_details(index, self.table_view)
        )
        self.table_last_games.doubleClicked.connect(
            lambda index: self.show_match_details(index, self.table_last_games)
        )
        self.table_future_games.doubleClicked.connect(
            lambda index: self.show_match_details(index, self.table_future_games)
        )

    def add_sale(self):
        """Открытие окна добавления продажи"""
        try:
            self.add_sale_window = AddSaleDialog(self.conn, self.cursor, self.user_role)
            self.add_sale_window.show()
            # Обновляем данные после закрытия окна
            self.add_sale_window.finished.connect(self.load_initial_data)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть окно добавления продажи:\n{e}")

    def add_inventory_item(self, medicine_id, batch_number, quantity, expiry_date, purchase_price):
        """Добавление товара в инвентарь с использованием точек сохранения"""
        try:
            with self.transaction_manager.transaction("Добавление в инвентарь") as tx_manager:
                # Точка сохранения: начало операции с инвентарем
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_START)
                
                # Проверяем существование лекарства
                self.cursor.execute("SELECT MedicineName FROM Medicine WHERE ID_Medicine = %s;", (medicine_id,))
                medicine_result = self.cursor.fetchone()
                if not medicine_result:
                    raise ValueError("Лекарство не найдено")
                
                medicine_name = medicine_result[0]
                
                # Точка сохранения: лекарство проверено
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_MEDICINE_CHECKED)
                
                # Проверяем срок годности
                if expiry_date <= datetime.date.today():
                    raise ValueError("Товар с истекшим сроком годности не может быть добавлен")
                
                # Добавляем товар в инвентарь
                self.cursor.execute("""
                    INSERT INTO Inventory (ID_Medicine, BatchNumber, Quantity, ExpiryDate, PurchasePrice, PurchaseDate)
                    VALUES (%s, %s, %s, %s, %s, CURRENT_DATE);
                """, (medicine_id, batch_number, quantity, expiry_date, purchase_price))
                
                # Точка сохранения: количество обновлено
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_QUANTITY_UPDATED)
                
                # Точка сохранения: операция завершена
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_COMPLETED)
            
            return True
            
        except Exception as e:
            print(f"Ошибка добавления в инвентарь: {e}")
            return False

    def update_inventory_quantity(self, inventory_id, new_quantity):
        """Обновление количества товара в инвентаре"""
        try:
            with self.transaction_manager.transaction("Обновление количества") as tx_manager:
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_START)
                
                # Проверяем существование записи
                self.cursor.execute("SELECT Quantity FROM Inventory WHERE ID_Inventory = %s;", (inventory_id,))
                result = self.cursor.fetchone()
                if not result:
                    raise ValueError("Запись инвентаря не найдена")
                
                old_quantity = result[0]
                
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_MEDICINE_CHECKED)
                
                # Обновляем количество
                self.cursor.execute("""
                    UPDATE Inventory 
                    SET Quantity = %s 
                    WHERE ID_Inventory = %s;
                """, (new_quantity, inventory_id))
                
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_QUANTITY_UPDATED)
                tx_manager.create_savepoint(PharmacySavePoints.INVENTORY_COMPLETED)
            
            return True
            
        except Exception as e:
            print(f"Ошибка обновления количества: {e}")
            return False

    def add_medicine(self, medicine_name, active_substance, dosage, form, prescription_required, price, supplier_id):
        """Добавление нового лекарства с использованием точек сохранения"""
        try:
            with self.transaction_manager.transaction("Добавление лекарства") as tx_manager:
                # Точка сохранения: начало операции с лекарством
                tx_manager.create_savepoint(PharmacySavePoints.MEDICINE_START)
                
                # Проверяем уникальность названия
                self.cursor.execute("SELECT ID_Medicine FROM Medicine WHERE MedicineName = %s;", (medicine_name,))
                if self.cursor.fetchone():
                    raise ValueError("Лекарство с таким названием уже существует")
                
                # Точка сохранения: лекарство валидировано
                tx_manager.create_savepoint(PharmacySavePoints.MEDICINE_VALIDATED)
                
                # Проверяем поставщика
                if supplier_id:
                    self.cursor.execute("SELECT SupplierName FROM Supplier WHERE ID_Supplier = %s;", (supplier_id,))
                    if not self.cursor.fetchone():
                        raise ValueError("Поставщик не найден")
                
                # Точка сохранения: поставщик проверен
                tx_manager.create_savepoint(PharmacySavePoints.MEDICINE_SUPPLIER_CHECKED)
                
                # Добавляем лекарство
                self.cursor.execute("""
                    INSERT INTO Medicine (MedicineName, ActiveSubstance, Dosage, Form, PrescriptionRequired, Price, ID_Supplier)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                    RETURNING ID_Medicine;
                """, (medicine_name, active_substance, dosage, form, prescription_required, price, supplier_id))
                
                medicine_id = self.cursor.fetchone()[0]
                
                # Точка сохранения: операция завершена
                tx_manager.create_savepoint(PharmacySavePoints.MEDICINE_COMPLETED)
            
            return medicine_id
            
        except Exception as e:
            print(f"Ошибка добавления лекарства: {e}")
            return None

    def add_supplier(self, supplier_name, contact_person, phone, email, address):
        """Добавление нового поставщика с использованием точек сохранения"""
        try:
            with self.transaction_manager.transaction("Добавление поставщика") as tx_manager:
                # Точка сохранения: начало операции с поставщиком
                tx_manager.create_savepoint(PharmacySavePoints.SUPPLIER_START)
                
                # Проверяем уникальность названия
                self.cursor.execute("SELECT ID_Supplier FROM Supplier WHERE SupplierName = %s;", (supplier_name,))
                if self.cursor.fetchone():
                    raise ValueError("Поставщик с таким названием уже существует")
                
                # Точка сохранения: поставщик валидирован
                tx_manager.create_savepoint(PharmacySavePoints.SUPPLIER_VALIDATED)
                
                # Добавляем поставщика
                self.cursor.execute("""
                    INSERT INTO Supplier (SupplierName, ContactPerson, Phone, Email, Address)
                    VALUES (%s, %s, %s, %s, %s)
                    RETURNING ID_Supplier;
                """, (supplier_name, contact_person, phone, email, address))
                
                supplier_id = self.cursor.fetchone()[0]
                
                # Точка сохранения: операция завершена
                tx_manager.create_savepoint(PharmacySavePoints.SUPPLIER_COMPLETED)
            
            return supplier_id
            
        except Exception as e:
            print(f"Ошибка добавления поставщика: {e}")
            return None

    def demonstrate_savepoints(self):
        """Демонстрация работы с точками сохранения"""
        try:
            print("=== Демонстрация системы точек сохранения ===")
            
            # Пример 1: Успешная операция с несколькими точками сохранения
            print("\n1. Успешная операция продажи:")
            with self.transaction_manager.transaction("Демо продажа") as tx_manager:
                sp1 = tx_manager.create_savepoint(PharmacySavePoints.SALE_START)
                print(f"   Создана точка сохранения: {sp1}")
                
                sp2 = tx_manager.create_savepoint(PharmacySavePoints.SALE_CUSTOMER_DATA)
                print(f"   Создана точка сохранения: {sp2}")
                
                sp3 = tx_manager.create_savepoint(PharmacySavePoints.SALE_ITEMS_ADDED)
                print(f"   Создана точка сохранения: {sp3}")
                
                sp4 = tx_manager.create_savepoint(PharmacySavePoints.SALE_COMPLETED)
                print(f"   Создана точка сохранения: {sp4}")
                
                print("   Транзакция завершена успешно!")
            
            # Пример 2: Операция с откатом к точке сохранения
            print("\n2. Операция с откатом:")
            try:
                with self.transaction_manager.transaction("Демо с ошибкой") as tx_manager:
                    sp1 = tx_manager.create_savepoint(PharmacySavePoints.MEDICINE_START)
                    print(f"   Создана точка сохранения: {sp1}")
                    
                    sp2 = tx_manager.create_savepoint(PharmacySavePoints.MEDICINE_VALIDATED)
                    print(f"   Создана точка сохранения: {sp2}")
                    
                    # Имитируем ошибку
                    raise ValueError("Имитация ошибки валидации")
                    
            except ValueError as e:
                print(f"   Ошибка перехвачена: {e}")
                print("   Откат к последней точке сохранения выполнен автоматически")
            
            # Пример 3: Показ текущих точек сохранения
            print(f"\n3. Текущие точки сохранения: {len(self.transaction_manager.savepoints)}")
            for i, sp in enumerate(self.transaction_manager.savepoints):
                print(f"   {i+1}. {sp}")
            
            print("\n=== Демонстрация завершена ===")
            
        except Exception as e:
            print(f"Ошибка демонстрации: {e}")

    def get_transaction_status(self):
        """Получение статуса текущей транзакции"""
        return {
            'active': self.transaction_manager.transaction_active,
            'savepoints': self.transaction_manager.savepoints.copy(),
            'count': len(self.transaction_manager.savepoints)
        }

    def test_registration(self):
        """Тест регистрации пользователя"""
        try:
            print("=== Тест регистрации ===")
            
            # Подключение к базе данных
            with psycopg2.connect(
                    dbname='pharmacydb',
                    user='pharmacy_admin',
                    password='pharmacyadmin',
                    host='localhost',
                    port=5432,
                    options='-c client_encoding=UTF8'
            ) as admin_conn:
                admin_cursor = admin_conn.cursor()
                
                # Проверяем структуру таблицы
                admin_cursor.execute("""
                    SELECT column_name, data_type, is_nullable 
                    FROM information_schema.columns 
                    WHERE table_name = 'dbusers' 
                    ORDER BY ordinal_position;
                """)
                
                columns = admin_cursor.fetchall()
                print("Структура таблицы dbusers:")
                for col in columns:
                    print(f"  {col[0]} - {col[1]} - {'NULL' if col[2] == 'YES' else 'NOT NULL'}")
                
                # Проверяем существующих пользователей
                admin_cursor.execute("SELECT COUNT(*) FROM dbusers;")
                count = admin_cursor.fetchone()[0]
                print(f"Количество пользователей в базе: {count}")
                
                # Пробуем создать тестового пользователя
                test_email = "test@pharmacy.com"
                test_password = "test123"
                
                admin_cursor.execute("SELECT 1 FROM dbusers WHERE user_email = %s", (test_email,))
                if admin_cursor.fetchone():
                    print(f"Пользователь {test_email} уже существует")
                else:
                    admin_cursor.execute("""
                        INSERT INTO dbusers (user_email, user_password, user_role)
                        VALUES (%s, %s, %s)
                    """, (test_email, test_password, 'pharmacy_user'))
                    admin_conn.commit()
                    print(f"Тестовый пользователь {test_email} создан успешно")
                
            print("=== Тест завершен ===")
            
        except Exception as e:
            print(f"Ошибка теста регистрации: {e}")
            import traceback
            traceback.print_exc()

    def show_sale_details(self, index, table_view):
        """Показ деталей продажи"""
        try:
            model = table_view.model()
            if not model or index.row() < 0:
                return
                
            # Получаем ID продажи из первой колонки
            sale_id = model.data(model.index(index.row(), 0))
            if not sale_id:
                return
                
            # Создаем и показываем окно деталей
            self.sale_details_window = SaleDetailsDialog(self.conn, self.cursor, sale_id)
            self.sale_details_window.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
            self.sale_details_window.show()
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть детали продажи:\n{e}")

    def show_inventory_details(self, index, table_view):
        """Показ деталей инвентаря"""
        try:
            model = table_view.model()
            if not model or index.row() < 0:
                return
                
            # Получаем данные о товаре
            medicine_name = model.data(model.index(index.row(), 0))
            batch_number = model.data(model.index(index.row(), 4))
            
            if not medicine_name:
                return
                
            # Показываем информацию о товаре
            QMessageBox.information(self, "Информация о товаре", 
                                  f"Лекарство: {medicine_name}\nПартия: {batch_number}")
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось показать детали инвентаря:\n{e}")

    def show_expiring_details(self, index, table_view):
        """Показ деталей товара с истекающим сроком"""
        try:
            model = table_view.model()
            if not model or index.row() < 0:
                return
                
            # Получаем данные о товаре
            medicine_name = model.data(model.index(index.row(), 0))
            expiry_date = model.data(model.index(index.row(), 6))
            status = model.data(model.index(index.row(), 8))
            
            if not medicine_name:
                return
                
            # Показываем предупреждение
            QMessageBox.warning(self, "Внимание! Товар с истекающим сроком", 
                              f"Лекарство: {medicine_name}\nСрок годности: {expiry_date}\nСтатус: {status}")
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось показать детали товара:\n{e}")

    def show_medicine_details(self, index, table_view):
        """Показ деталей лекарства"""
        try:
            model = table_view.model()
            if not model or index.row() < 0:
                return
                
            # Получаем данные о лекарстве
            medicine_name = model.data(model.index(index.row(), 0))
            active_substance = model.data(model.index(index.row(), 1))
            dosage = model.data(model.index(index.row(), 2))
            price = model.data(model.index(index.row(), 5))
            quantity = model.data(model.index(index.row(), 7))
            
            if not medicine_name:
                return
                
            # Показываем информацию о лекарстве
            QMessageBox.information(self, "Информация о лекарстве", 
                                  f"Название: {medicine_name}\n"
                                  f"Действующее вещество: {active_substance}\n"
                                  f"Дозировка: {dosage}\n"
                                  f"Цена: {price}\n"
                                  f"На складе: {quantity}")
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось показать детали лекарства:\n{e}")

    def show_match_details(self, index, table_view):
        """Безопасное открытие окна с деталями матча"""
        try:
            model = table_view.model()
            if not hasattr(model, '_table_data'):
                return

            row = index.row()
            if row < 0 or row >= len(model._table_data):
                return

            row_data = model._table_data[row]
            if len(row_data) < 7:
                raise ValueError("Недостаточно данных в строке таблицы")

            # Создаем словарь с данными матча
            match_data = {
                'date': row_data[0],
                'team1': row_data[1],
                'team2': row_data[2],
                'score1': row_data[3],
                'score2': row_data[4],
                'sport': row_data[5],
                'tournament': row_data[6]
            }

            # Создаем и показываем окно деталей
            self.detail_window = MatchDetail(self.conn, self.cursor, match_data)
            self.detail_window.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
            self.detail_window.show()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка",
                                 f"Не удалось открыть детали матча:\n{str(e)[:200]}")
            print(f"Ошибка: {e}")
            import traceback
            traceback.print_exc()


class GPTWorker(QObject):
    analysis_done = pyqtSignal(str)
    error_occurred = pyqtSignal(str)
    finished = pyqtSignal()

    def __init__(self, team_name):
        super().__init__()
        self.team_name = team_name
        self.max_retries = 2
        self._should_stop = False
        self.api_url = "https://api.deepseek.com/v1/chat/completions"
        self.headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {os.getenv('sk-6716f9ed1a934a7bb6a996ef75633deb')}"
        }

    @pyqtSlot()
    def run_analysis(self):
        result = None
        error = None

        for attempt in range(self.max_retries):
            if self._should_stop:
                break

            try:
                payload = {
                    "model": "deepseek-chat",
                    "messages": [{
                        "role": "user",
                        "content": f"Сделай профессиональный анализ баскетбольной команды {self.team_name}. Опиши их стиль игры, сильные и слабые стороны, ключевых игроков. Ответ ограничь 150 словами."
                    }],
                    "temperature": 0.7,
                    "max_tokens": 300
                }

                response = requests.post(
                    self.api_url,
                    headers=self.headers,
                    json=payload,
                    timeout=15
                )

                # Проверка статуса ответа
                if response.status_code != 200:
                    raise ValueError(f"API error {response.status_code}: {response.text[:200]}")

                data = response.json()
                if "choices" not in data or not data["choices"]:
                    raise ValueError("Invalid response format from API")

                content = data["choices"][0]["message"]["content"]
                result = content[:1000]  # Обрезаем слишком длинный ответ
                break

            except requests.exceptions.RequestException as e:
                error = f"Сетевая ошибка: {str(e)[:150]}"
            except ValueError as e:
                error = f"Ошибка API: {str(e)[:150]}"
            except Exception as e:
                error = f"Неожиданная ошибка: {str(e)[:150]}"

            print(f"Попытка {attempt + 1} провалена: {error}")
            time.sleep(1)  # Задержка между попытками

        if result:
            self.analysis_done.emit(result)
        else:
            self.error_occurred.emit(error or "Не удалось получить анализ после всех попыток")

        self.finished.emit()

    def stop(self):
        self._should_stop = True


class MatchDetail(QDialog):
    def __init__(self, conn, cursor, match_data):
        super().__init__()
        self.conn = conn
        self.cursor = cursor
        self.match_data = match_data
        self._init_workers()
        self.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
        self.setup_ui()

    def _init_workers(self):
        """Initialize worker-related attributes"""
        self.gpt_thread1 = self.gpt_thread2 = None
        self.gpt_worker1 = self.gpt_worker2 = None

    def setup_ui(self):
        """Initialize UI components safely"""
        try:
            if not hasattr(self.match_data, 'team1') or not hasattr(self.match_data, 'team2'):
                raise ValueError("Invalid match_data: missing team names")

            ui_path = resource_path("details.ui")
            if not os.path.exists(ui_path):
                raise FileNotFoundError(f"UI file not found: {ui_path}")

            loadUi(ui_path, self)
            self.setWindowTitle(f"Match details: {self.match_data.team1} vs {self.match_data.team2}")

            # Initialize widgets with type hints
            self.list_analytic_one: QTextEdit = self.findChild(QTextEdit, "textEdit") or QTextEdit(self)
            self.list_analytic_two: QTextEdit = self.findChild(QTextEdit, "textEdit_2") or QTextEdit(self)

            self.table_teams_names: QTableView = self.findChild(QTableView, "tableView_4") or QTableView(self)
            self.table_team_one: QTableView = self.findChild(QTableView, "tableView_5") or QTableView(self)
            self.table_team_two: QTableView = self.findChild(QTableView, "tableView_6") or QTableView(self)

            # Clear initial content
            for widget in [self.list_analytic_one, self.list_analytic_two]:
                widget.setPlainText("")

            # Layout optimization
            if layout := self.layout():
                layout.setSizeConstraint(QLayout.SetFixedSize)

            QTimer.singleShot(100, self.load_all_data)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"UI initialization failed:\n{str(e)[:200]}")
            self.close()

    def load_all_data(self):
        """Load all data components with safety checks"""
        try:
            if not self._check_connection():
                QMessageBox.critical(self, "Error", "Database connection failed")
                self.close()
                return

            with self._db_operation():
                self.cleanup_models()
                self.setup_teams_table()
                self.load_team_players(self.match_data.team1, self.table_team_one)
                self.load_team_players(self.match_data.team2, self.table_team_two)
                self.load_match_stats()
                self.load_analytics()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Data loading failed:\n{str(e)[:200]}")

    @contextmanager
    def _db_operation(self):
        """Context manager for database operations"""
        try:
            yield
        except Exception as e:
            self.conn.rollback()
            raise
        finally:
            self.conn.commit()

    def load_analytics(self):
        """Initialize and start analysis threads"""
        try:
            # Set loading state
            self._set_loading_state(True)

            # Configure and start threads
            self.gpt_thread1 = self._create_worker_thread(
                team=self.match_data.team1,
                result_handler=self.on_team1_result,
                error_handler=self.on_team1_error
            )

            self.gpt_thread2 = self._create_worker_thread(
                team=self.match_data.team2,
                result_handler=self.on_team2_result,
                error_handler=self.on_team2_error
            )

            self.gpt_thread1.start()
            self.gpt_thread2.start()

        except Exception as e:
            error_msg = f"Analysis initialization failed: {str(e)}"
            print(error_msg)
            self._set_loading_state(False, error_msg)

    def _create_worker_thread(self, team, result_handler, error_handler):
        """Helper to create configured worker thread"""
        thread = QThread()
        worker = GPTWorker(team)
        worker.moveToThread(thread)

        # Connect signals
        worker.analysis_done.connect(result_handler)
        worker.error_occurred.connect(error_handler)
        worker.finished.connect(thread.quit)
        thread.started.connect(worker.run_analysis)

        # Cleanup connections
        thread.finished.connect(thread.deleteLater)
        thread.finished.connect(worker.deleteLater)

        return thread

    def _set_loading_state(self, loading, error_msg=None):
        """Update UI loading state"""
        if loading:
            text = "Generating analysis..."
            status = text
        else:
            text = error_msg if error_msg else "Analysis complete"
            status = "Ready"

        for widget, team in [
            (self.list_analytic_one, self.match_data.team1),
            (self.list_analytic_two, self.match_data.team2)
        ]:
            if loading:
                widget.setPlainText(f"Generating analysis for {team}...")
            elif error_msg:
                widget.setPlainText(f"Error for {team}: {error_msg}")

    @pyqtSlot(str)
    def on_team1_result(self, text):
        self._safe_set_text(self.list_analytic_one, text)

    @pyqtSlot(str)
    def on_team2_result(self, text):
        self._safe_set_text(self.list_analytic_two, text)

    @pyqtSlot(str)
    def on_team1_error(self, error):
        print(f"Team 1 error: {error}")
        self._safe_set_text(self.list_analytic_one, error)

    @pyqtSlot(str)
    def on_team2_error(self, error):
        print(f"Team 2 error: {error}")
        self._safe_set_text(self.list_analytic_two, error)

    def _safe_set_text(self, widget, text):
        """Thread-safe text setting"""
        if widget and isinstance(widget, QTextEdit):
            widget.setPlainText(str(text)[:1000])  # Limit text length

    def cleanup_models(self):
        """Clean up table models safely"""
        for table in [self.table_teams_names, self.table_team_one, self.table_team_two]:
            if table and isinstance(table, QTableView):
                if model := table.model():
                    model.deleteLater()
                table.setModel(None)

    def closeEvent(self, event):
        """Handle window closing with proper cleanup"""
        print("Closing window, performing cleanup...")

        # Stop and cleanup workers
        for worker, thread in zip(
                [self.gpt_worker1, self.gpt_worker2],
                [self.gpt_thread1, self.gpt_thread2]
        ):
            if worker:
                worker.stop()
                worker.disconnect()

            if thread and thread.isRunning():
                thread.quit()
                if not thread.wait(1000):
                    thread.terminate()

        # Cleanup models and references
        self.cleanup_models()
        self._init_workers()  # Reset all worker references

        # Ensure database connection is clean
        try:
            self.conn.commit()
        except:
            pass

        event.accept()

    def _check_connection(self):
        """Verify database connection"""
        try:
            self.cursor.execute("SELECT 1")
            return True
        except Exception as e:
            print(f"Database connection check failed: {e}")
            return False

    def setup_teams_table(self):
        """Initialize teams comparison table with scores"""
        try:
            # Создаем модель с 3 колонками: Team 1, Score, Team 2
            model = QStandardItemModel(1, 3)
            model.setHorizontalHeaderLabels(["Team 1", "Score", "Team 2"])

            # Добавляем названия команд
            team1_item = QStandardItem(self.match_data.team1)
            team1_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)

            score_item = QStandardItem(f"{self.match_data.score1} - {self.match_data.score2}")
            score_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)

            team2_item = QStandardItem(self.match_data.team2)
            team2_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)

            # Устанавливаем элементы в модель
            model.setItem(0, 0, team1_item)
            model.setItem(0, 1, score_item)
            model.setItem(0, 2, team2_item)

            # Настраиваем таблицу
            self.table_teams_names.setModel(model)
            header = self.table_teams_names.horizontalHeader()

            # Универсальная настройка растягивания колонок
            header.setSectionResizeMode(QHeaderView.ResizeMode.Stretch)


        except Exception as e:
            print(f"Teams table setup failed: {e}")
            QMessageBox.warning(self, "Error", f"Failed to setup teams table: {str(e)[:200]}")

    def load_team_players(self, team_name, table_view):
        """Load players for specified team"""
        try:
            with self._db_operation():
                self.cursor.execute("""
                    SELECT FirstName, LastName, Position, DateOfBirth
                    FROM Player p JOIN Team t ON p.ID_Team = t.ID_Team
                    WHERE t.TeamName = %s
                """, (team_name,))

                model = QStandardItemModel()
                model.setHorizontalHeaderLabels(
                    ["First Name", "Last Name", "Position", "Birth Date"]
                )

                for row in self.cursor.fetchall():
                    model.appendRow([QStandardItem(str(x or "")) for x in row])

                table_view.setModel(model)
                table_view.resizeColumnsToContents()

        except Exception as e:
            print(f"Failed to load players for {team_name}: {e}")
            QMessageBox.warning(
                self,
                "Error",
                f"Failed to load players for {team_name}:\n{str(e)[:200]}"
            )

    def load_match_stats(self):
        """Load match statistics - to be implemented"""
        try:
            # Implementation goes here
            pass
        except Exception as e:
            print(f"Match stats loading failed: {e}")


def resource_path(relative_path):
    """Возвращает абсолютный путь к ресурсу"""
    try:
        base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
        return os.path.join(base_path, relative_path)
    except Exception:
        return relative_path


class AddSaleDialog(QDialog):
    def __init__(self, conn, cursor, user_role):
        super().__init__()
        self.conn = conn
        self.cursor = cursor
        self.user_role = user_role
        self.sale_items = []  # Список товаров в продаже
        
        loadUi(resource_path("add_sale.ui"), self)
        self.setFixedSize(self.size())
        
        # Подключение виджетов
        self.customer_name_edit = self.findChild(QLineEdit, "lineEdit_customer_name")
        self.customer_phone_edit = self.findChild(QLineEdit, "lineEdit_customer_phone")
        self.payment_method_combo = self.findChild(QComboBox, "comboBox_payment_method")
        self.medicine_combo = self.findChild(QComboBox, "comboBox_medicine")
        self.quantity_spinbox = self.findChild(QSpinBox, "spinBox_quantity")
        self.price_edit = self.findChild(QLineEdit, "lineEdit_price")
        self.add_item_button = self.findChild(QPushButton, "pushButton_add_item")
        self.items_table = self.findChild(QTableWidget, "tableWidget_items")
        self.total_label = self.findChild(QLabel, "label_total_value")
        self.save_button = self.findChild(QPushButton, "pushButton_save")
        self.cancel_button = self.findChild(QPushButton, "pushButton_cancel")
        
        # Подключение сигналов
        self.add_item_button.clicked.connect(self.add_item_to_sale)
        self.save_button.clicked.connect(self.save_sale)
        self.cancel_button.clicked.connect(self.reject)
        self.medicine_combo.currentTextChanged.connect(self.update_price)
        
        # Загрузка данных
        self.load_medicines()
        self.update_total()
        
    def load_medicines(self):
        """Загрузка списка лекарств"""
        try:
            self.cursor.execute("SELECT MedicineName, Price FROM Medicine ORDER BY MedicineName;")
            medicines = self.cursor.fetchall()
            self.medicine_combo.clear()
            for medicine_name, price in medicines:
                self.medicine_combo.addItem(medicine_name)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки лекарств:\n{e}")
    
    def update_price(self):
        """Обновление цены при выборе лекарства"""
        try:
            medicine_name = self.medicine_combo.currentText()
            if medicine_name:
                self.cursor.execute("SELECT Price FROM Medicine WHERE MedicineName = %s;", (medicine_name,))
                result = self.cursor.fetchone()
                if result:
                    self.price_edit.setText(f"{result[0]:.2f}")
                else:
                    self.price_edit.setText("0.00")
        except Exception as e:
            self.price_edit.setText("0.00")
    
    def add_item_to_sale(self):
        """Добавление товара в продажу"""
        try:
            medicine_name = self.medicine_combo.currentText()
            quantity = self.quantity_spinbox.value()
            price_text = self.price_edit.text()
            
            if not medicine_name:
                QMessageBox.warning(self, "Ошибка", "Выберите лекарство")
                return
                
            if not price_text or float(price_text) <= 0:
                QMessageBox.warning(self, "Ошибка", "Неверная цена")
                return
            
            price = float(price_text)
            total_price = price * quantity
            
            # Проверяем наличие на складе
            self.cursor.execute("""
                SELECT COALESCE(SUM(i.Quantity), 0) 
                FROM Inventory i 
                JOIN Medicine m ON i.ID_Medicine = m.ID_Medicine 
                WHERE m.MedicineName = %s AND i.ExpiryDate > CURRENT_DATE
            """, (medicine_name,))
            
            available_quantity = self.cursor.fetchone()[0]
            if available_quantity < quantity:
                QMessageBox.warning(self, "Ошибка", f"Недостаточно товара на складе. Доступно: {available_quantity}")
                return
            
            # Добавляем товар в список
            item = {
                'medicine_name': medicine_name,
                'quantity': quantity,
                'price': price,
                'total_price': total_price
            }
            self.sale_items.append(item)
            
            # Обновляем таблицу
            self.update_items_table()
            self.update_total()
            
            # Очищаем поля
            self.quantity_spinbox.setValue(1)
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка добавления товара:\n{e}")
    
    def update_items_table(self):
        """Обновление таблицы товаров"""
        self.items_table.setRowCount(len(self.sale_items))
        
        for row, item in enumerate(self.sale_items):
            self.items_table.setItem(row, 0, QTableWidgetItem(item['medicine_name']))
            self.items_table.setItem(row, 1, QTableWidgetItem(str(item['quantity'])))
            self.items_table.setItem(row, 2, QTableWidgetItem(f"{item['price']:.2f}"))
            self.items_table.setItem(row, 3, QTableWidgetItem(f"{item['total_price']:.2f}"))
            
            # Кнопка удаления
            remove_button = QPushButton("Удалить")
            remove_button.clicked.connect(lambda checked, r=row: self.remove_item(r))
            self.items_table.setCellWidget(row, 4, remove_button)
    
    def remove_item(self, row):
        """Удаление товара из продажи"""
        if 0 <= row < len(self.sale_items):
            self.sale_items.pop(row)
            self.update_items_table()
            self.update_total()
    
    def update_total(self):
        """Обновление общей суммы"""
        total = sum(item['total_price'] for item in self.sale_items)
        self.total_label.setText(f"{total:.2f} руб.")
    
    def save_sale(self):
        """Сохранение продажи с использованием точек сохранения"""
        try:
            if not self.sale_items:
                QMessageBox.warning(self, "Ошибка", "Добавьте товары в продажу")
                return
            
            customer_name = self.customer_name_edit.text().strip() or "Не указано"
            customer_phone = self.customer_phone_edit.text().strip() or "Не указано"
            payment_method = self.payment_method_combo.currentText()
            total_amount = sum(item['total_price'] for item in self.sale_items)
            
            # Получаем ID пользователя
            user_id = 1  # По умолчанию, можно получить из сессии
            
            # Используем менеджер транзакций с точками сохранения
            with self.transaction_manager.transaction("Создание продажи") as tx_manager:
                # Точка сохранения: начало операции продажи
                tx_manager.create_savepoint(PharmacySavePoints.SALE_START)
                
                # Создаем продажу
                self.cursor.execute("""
                    INSERT INTO Sales (CustomerName, CustomerPhone, TotalAmount, PaymentMethod, ID_User)
                    VALUES (%s, %s, %s, %s, %s)
                    RETURNING ID_Sale;
                """, (customer_name, customer_phone, total_amount, payment_method, user_id))
                
                sale_id = self.cursor.fetchone()[0]
                
                # Точка сохранения: данные покупателя сохранены
                tx_manager.create_savepoint(PharmacySavePoints.SALE_CUSTOMER_DATA)
                
                # Добавляем товары с проверкой остатков
                for i, item in enumerate(self.sale_items):
                    try:
                        # Получаем ID лекарства
                        self.cursor.execute("SELECT ID_Medicine FROM Medicine WHERE MedicineName = %s;", (item['medicine_name'],))
                        medicine_id = self.cursor.fetchone()[0]
                        
                        # Проверяем остатки на складе
                        self.cursor.execute("""
                            SELECT COALESCE(SUM(i.Quantity), 0) 
                            FROM Inventory i 
                            WHERE i.ID_Medicine = %s AND i.ExpiryDate > CURRENT_DATE
                        """, (medicine_id,))
                        
                        available_quantity = self.cursor.fetchone()[0]
                        if available_quantity < item['quantity']:
                            raise ValueError(f"Недостаточно товара '{item['medicine_name']}' на складе. Доступно: {available_quantity}")
                        
                        # Добавляем позицию продажи
                        self.cursor.execute("""
                            INSERT INTO SaleItems (ID_Sale, ID_Medicine, Quantity, UnitPrice, TotalPrice)
                            VALUES (%s, %s, %s, %s, %s);
                        """, (sale_id, medicine_id, item['quantity'], item['price'], item['total_price']))
                        
                        # Точка сохранения после каждого товара
                        if i == 0:  # Первый товар
                            tx_manager.create_savepoint(PharmacySavePoints.SALE_ITEMS_ADDED)
                        
                    except Exception as e:
                        # Откат к точке сохранения перед добавлением товаров
                        tx_manager.rollback_to_savepoint(PharmacySavePoints.SALE_CUSTOMER_DATA)
                        raise ValueError(f"Ошибка добавления товара '{item['medicine_name']}': {e}")
                
                # Точка сохранения: товары добавлены
                tx_manager.create_savepoint(PharmacySavePoints.SALE_INVENTORY_UPDATED)
                
                # Триггеры автоматически обновят остатки на складе
                # Точка сохранения: операция завершена
                tx_manager.create_savepoint(PharmacySavePoints.SALE_COMPLETED)
            
            QMessageBox.information(self, "Успех", f"Продажа #{sale_id} сохранена успешно!")
            self.accept()
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка сохранения продажи:\n{e}")
            print(f"Детали ошибки: {e}")
            import traceback
            traceback.print_exc()


class SaleDetailsDialog(QDialog):
    def __init__(self, conn, cursor, sale_id):
        super().__init__()
        self.conn = conn
        self.cursor = cursor
        self.sale_id = sale_id
        
        loadUi(resource_path("sale_details.ui"), self)
        self.setFixedSize(self.size())
        
        # Подключение виджетов
        self.sale_id_label = self.findChild(QLabel, "label_sale_id_value")
        self.sale_date_label = self.findChild(QLabel, "label_sale_date_value")
        self.customer_name_label = self.findChild(QLabel, "label_customer_name_value")
        self.customer_phone_label = self.findChild(QLabel, "label_customer_phone_value")
        self.total_amount_label = self.findChild(QLabel, "label_total_amount_value")
        self.payment_method_label = self.findChild(QLabel, "label_payment_method_value")
        self.items_table = self.findChild(QTableWidget, "tableWidget_items")
        self.print_button = self.findChild(QPushButton, "pushButton_print")
        self.close_button = self.findChild(QPushButton, "pushButton_close")
        
        # Подключение сигналов
        self.print_button.clicked.connect(self.print_receipt)
        self.close_button.clicked.connect(self.accept)
        
        # Загрузка данных
        self.load_sale_data()
    
    def load_sale_data(self):
        """Загрузка данных о продаже"""
        try:
            # Загружаем основную информацию о продаже
            self.cursor.execute("""
                SELECT SaleDate, CustomerName, CustomerPhone, TotalAmount, PaymentMethod
                FROM Sales WHERE ID_Sale = %s;
            """, (self.sale_id,))
            
            sale_data = self.cursor.fetchone()
            if not sale_data:
                QMessageBox.critical(self, "Ошибка", "Продажа не найдена")
                self.reject()
                return
            
            # Заполняем поля
            self.sale_id_label.setText(str(self.sale_id))
            self.sale_date_label.setText(sale_data[0].strftime("%Y-%m-%d %H:%M"))
            self.customer_name_label.setText(sale_data[1] or "Не указано")
            self.customer_phone_label.setText(sale_data[2] or "Не указано")
            self.total_amount_label.setText(f"{sale_data[3]:.2f} руб.")
            self.payment_method_label.setText(sale_data[4])
            
            # Загружаем товары
            self.cursor.execute("""
                SELECT m.MedicineName, m.ActiveSubstance, m.Dosage, m.Form,
                       si.Quantity, si.UnitPrice, si.TotalPrice
                FROM SaleItems si
                JOIN Medicine m ON si.ID_Medicine = m.ID_Medicine
                WHERE si.ID_Sale = %s;
            """, (self.sale_id,))
            
            items_data = self.cursor.fetchall()
            
            # Заполняем таблицу товаров
            self.items_table.setRowCount(len(items_data))
            for row, item in enumerate(items_data):
                self.items_table.setItem(row, 0, QTableWidgetItem(item[0]))  # Лекарство
                self.items_table.setItem(row, 1, QTableWidgetItem(item[1] or "Не указано"))  # Действующее вещество
                self.items_table.setItem(row, 2, QTableWidgetItem(item[2] or "Не указано"))  # Дозировка
                self.items_table.setItem(row, 3, QTableWidgetItem(item[3] or "Не указано"))  # Форма
                self.items_table.setItem(row, 4, QTableWidgetItem(str(item[4])))  # Количество
                self.items_table.setItem(row, 5, QTableWidgetItem(f"{item[5]:.2f}"))  # Цена за ед.
                self.items_table.setItem(row, 6, QTableWidgetItem(f"{item[6]:.2f}"))  # Сумма
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки данных продажи:\n{e}")
    
    def print_receipt(self):
        """Печать чека"""
        QMessageBox.information(self, "Печать", "Функция печати чека будет реализована в следующей версии")


class AddResult(QDialog):
    def __init__(self, conn, cursor):
        super().__init__()
        if not hasattr(cursor, "execute"):
            raise ValueError("Ошибка: передан неправильный объект вместо курсора базы данных.")

        self.conn = conn
        self.cursor = cursor
        loadUi(resource_path("addmatch.ui"), self)
        self.setFixedSize(self.size())

        # Подключаем виджеты из нового UI
        self.sport_combo = self.findChild(QComboBox, "comboBox")
        self.tournament_combo = self.findChild(QComboBox, "comboBox_2")
        self.select_team1 = self.findChild(QComboBox, "comboBox_3")
        self.select_team2 = self.findChild(QComboBox, "comboBox_4")
        self.check_team1 = self.findChild(QCheckBox, "checkBox")
        self.check_team2 = self.findChild(QCheckBox, "checkBox_2")
        self.date_edit = self.findChild(QDateEdit, "dateEdit")
        self.team1_line = self.findChild(QLineEdit, "lineEdit")
        self.team2_line = self.findChild(QLineEdit, "lineEdit_2")
        self.score_line = self.findChild(QLineEdit, "lineEdit_3")
        self.add_button = self.findChild(QPushButton, "pushButton_2")
        self.add_participants_button = self.findChild(QPushButton, "pushButton_3")

        # Устанавливаем текущую дату по умолчанию
        self.date_edit.setDate(QDate.currentDate())

        # Настройка видимости полей ввода
        self.team1_line.setVisible(False)
        self.team2_line.setVisible(False)
        self.check_team1.setChecked(False)
        self.check_team2.setChecked(False)

        # Подключаем сигналы чекбоксов
        self.check_team1.stateChanged.connect(lambda: self.team1_line.setVisible(self.check_team1.isChecked()))
        self.check_team2.stateChanged.connect(lambda: self.team2_line.setVisible(self.check_team2.isChecked()))

        # Подключаем сигналы кнопок
        self.add_button.clicked.connect(self.handle_add_button_click)
        self.add_participants_button.clicked.connect(self.handle_add_participants_click)

        self.load_sports()
        self.load_tournaments()
        self.load_teams()
        self.load_teams1()

    @with_sound("click")
    def handle_add_button_click(self):
        """Обработчик клика по кнопке добавления матча"""
        self.save_result()

    @with_sound("click")
    def handle_add_participants_click(self):
        """Обработчик клика по кнопке добавления участников"""
        self.add_participants()

    def load_teams(self):
        try:
            self.cursor.execute("SELECT TeamName FROM Team;")
            teams = self.cursor.fetchall()
            self.select_team1.clear()
            self.select_team1.addItem("Все команды")
            for team in teams:
                self.select_team1.addItem(team[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки команд:\n{e}")

    def load_teams1(self):
        try:
            self.cursor.execute("SELECT TeamName FROM Team;")
            teams = self.cursor.fetchall()
            self.select_team2.clear()
            self.select_team2.addItem("Все команды")
            for team in teams:
                self.select_team2.addItem(team[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки команд:\n{e}")

    def add_participants(self):
        try:
            team1_name = self.select_team1.currentText()
            team2_name = self.select_team2.currentText()

            if team1_name in ("Все команды", "") or team2_name in ("Все команды", ""):
                QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите конкретные команды для обеих сторон.")
                return

            sport_id = self.get_sport_id(self.sport_combo.currentText())
            if not sport_id:
                QMessageBox.warning(self, "Ошибка", "Не выбран вид спорта")
                return

            self.cursor.execute(
                "SELECT ID_Team FROM Team WHERE TeamName = %s AND ID_SportType = %s;",
                (team1_name, sport_id)
            )
            row = self.cursor.fetchone()
            if not row:
                QMessageBox.warning(self, "Ошибка", f"Команда {team1_name} не найдена")
                return
            team1_id = row[0]

            self.cursor.execute(
                "SELECT ID_Team FROM Team WHERE TeamName = %s AND ID_SportType = %s;",
                (team2_name, sport_id)
            )
            row = self.cursor.fetchone()
            if not row:
                QMessageBox.warning(self, "Ошибка", f"Команда {team2_name} не найдена")
                return
            team2_id = row[0]

            # предотвращаем двойной запуск
            self.add_participants_button.setEnabled(False)
            try:
                # Делаем диалог дочерним и храним ссылку
                self.players_dialog = AddPlayersWindow(
                    parent=self,  # ВАЖНО: родитель
                    conn=self.conn,
                    cursor=self.cursor,
                    team1_id=team1_id,
                    team2_id=team2_id,
                    team1_name=team1_name,
                    team2_name=team2_name
                )
                self.players_dialog.setModal(True)
                self.players_dialog.exec()
            finally:
                self.add_participants_button.setEnabled(True)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть окно добавления игроков:\n{str(e)}")
            import traceback;
            print(traceback.format_exc())

    def get_or_create_team(self, team_name, sport_id):
        """Получаем или создаем команду с учетом sport_id"""
        self.cursor.execute(
            "SELECT ID_Team FROM Team WHERE TeamName = %s AND ID_SportType = %s;",
            (team_name, sport_id)
        )
        result = self.cursor.fetchone()
        if result:
            return result[0]

        self.cursor.execute(
            "INSERT INTO Team (TeamName, ID_SportType) VALUES (%s, %s) RETURNING ID_Team;",
            (team_name, sport_id)
        )
        return self.cursor.fetchone()[0]

    def load_sports(self):
        try:
            self.cursor.execute("SELECT SportName FROM SportType;")
            sports = self.cursor.fetchall()
            self.sport_combo.clear()
            for sport in sports:
                self.sport_combo.addItem(sport[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки видов спорта:\n{e}")
            print(f"Ошибка загрузки видов спорта: {e}")

    def load_tournaments(self):
        try:
            self.cursor.execute("SELECT TournamentName FROM Tournament;")
            tournaments = self.cursor.fetchall()
            self.tournament_combo.clear()
            self.tournament_combo.addItem("Без турнира")
            for tournament in tournaments:
                self.tournament_combo.addItem(tournament[0])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки турниров:\n{e}")
            print(f"Ошибка загрузки турниров: {e}")

    def save_result(self):
        try:
            data = self.validate_and_get_input()
            if not data:
                return

            sport_name, tournament_name, match_date, team1_name, team2_name, score1, score2 = data

            print(f"Попытка добавить матч: {team1_name} vs {team2_name}, счет {score1}:{score2}")

            # Получаем ID спорта
            sport_id = self.get_sport_id(sport_name)
            if sport_id is None:
                QMessageBox.critical(self, "Ошибка", "Вид спорта не найден.")
                return
            print(f"ID спорта: {sport_id}")

            # Получаем ID турнира (может быть None)
            tournament_id = None
            if tournament_name != "Без турнира":
                tournament_id = self.get_tournament_id(tournament_name)
                if tournament_id is None:
                    QMessageBox.critical(self, "Ошибка", "Турнир не найден.")
                    return
            print(f"ID турнира: {tournament_id}")

            # Получаем или создаем команды
            team1_id = self.get_or_create_team(team1_name, sport_id)
            team2_id = self.get_or_create_team(team2_name, sport_id)
            print(f"ID команд: {team1_id} и {team2_id}")

            # Добавляем матч (всегда создаем новую запись)
            match_id = self.insert_match(match_date, team1_id, team2_id, sport_id, tournament_id)
            print(f"ID матча: {match_id}")

            # Добавляем результаты (всегда создаем новую запись)
            self.insert_result(match_id, team1_id, score1, team2_id, score2)
            print("Результаты добавлены")

            # Используем новую систему точек сохранения
            transaction_manager = PharmacyTransactionManager(self.conn, self.cursor)
            with transaction_manager.transaction("Сохранение результата матча") as tx_manager:
                tx_manager.create_savepoint("match_validation")
                tx_manager.create_savepoint("match_created")
                tx_manager.create_savepoint("results_added")
                tx_manager.create_savepoint("match_completed")
            
            QMessageBox.information(self, "Успешно", "Матч и результат добавлены!")
            print("Данные успешно сохранены")

            # Очищаем поля
            self.team1_line.clear()
            self.team2_line.clear()
            self.score_line.clear()

        except Exception as e:
            # Используем новую систему отката
            transaction_manager = PharmacyTransactionManager(self.conn, self.cursor)
            transaction_manager.rollback_all()
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить результат:\n{e}")
            print(f"Ошибка при сохранении: {e}")
            import traceback
            traceback.print_exc()

    def validate_and_get_input(self):
        sport_name = self.sport_combo.currentText()
        tournament_name = self.tournament_combo.currentText()
        match_date = self.date_edit.date().toString("yyyy-MM-dd")
        team1_name = self.team1_line.text().strip()
        team2_name = self.team2_line.text().strip()
        score_text = self.score_line.text().strip()

        # Проверяем, что команды заполнены
        if not team1_name or not team2_name:
            QMessageBox.warning(self, "Ошибка", "Необходимо указать обе команды для матча.")
            return None

        if not score_text:
            QMessageBox.warning(self, "Ошибка", "Введите счет матча.")
            return None

        try:
            score1, score2 = map(int, score_text.split())
        except ValueError:
            QMessageBox.warning(self, "Неверный формат", "Введите счёт в формате: '1 2' (два числа через пробел).")
            return None

        return sport_name, tournament_name, match_date, team1_name, team2_name, score1, score2

    def get_sport_id(self, sport_name):
        self.cursor.execute("SELECT ID_SportType FROM SportType WHERE SportName = %s;", (sport_name,))
        result = self.cursor.fetchone()
        return result[0] if result else None

    def get_tournament_id(self, name):
        self.cursor.execute("SELECT ID_Tournament FROM Tournament WHERE TournamentName = %s;", (name,))
        result = self.cursor.fetchone()
        return result[0] if result else None

    def insert_match(self, match_date, team1_id, team2_id, sport_id, tournament_id):
        query = """
            INSERT INTO Match (MatchDateTime, ID_Team1, ID_Team2, ID_SportType, ID_Tournament)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING ID_Match;
        """
        params = (match_date, team1_id, team2_id, sport_id, tournament_id)

        print(f"Выполняем запрос: {query % params}")

        self.cursor.execute(query, params)
        return self.cursor.fetchone()[0]

    def insert_result(self, match_id, team1_id, score1, team2_id, score2):
        # Проверяем, существуют ли уже результаты для этого матча
        self.cursor.execute("SELECT COUNT(*) FROM Result WHERE ID_Match = %s;", (match_id,))
        if self.cursor.fetchone()[0] > 0:
            print(f"Результаты для матча {match_id} уже существуют. Обновляем.")
            # Можно добавить логику обновления существующих результатов
            return

        query = """
            INSERT INTO Result (ID_Match, ID_Team, Score)
            VALUES (%s, %s, %s), (%s, %s, %s);
        """
        params = (match_id, team1_id, score1, match_id, team2_id, score2)

        print(f"Выполняем запрос: {query % params}")
        self.cursor.execute(query, params)


class AddPlayersWindow(QDialog):
    def __init__(self, parent, conn, cursor, team1_id, team2_id, team1_name, team2_name):
        super().__init__(parent)
        self.conn = conn
        self.cursor = cursor
        self.team1_id = team1_id
        self.team2_id = team2_id

        self._load_ui()
        self._init_widgets(team1_name, team2_name)
        self._setup_tables()
        self._load_existing_players()
        self.pushButton_3.clicked.connect(self._save_players)

    def _load_ui(self):
        ui_file = resource_path("addteam.ui")
        if not os.path.exists(ui_file):
            QMessageBox.critical(self, "Ошибка", f"UI файл не найден: {ui_file}")
            raise FileNotFoundError(f"UI файл не найден: {ui_file}")
        loadUi(ui_file, self)

    def _ensure_item(self, table, row, col):
        item = table.item(row, col)
        if item is None:
            item = QTableWidgetItem("")
            table.setItem(row, col, item)
        return item


    def _init_widgets(self, team1_name, team2_name):
        self.label_2 = self.findChild(QLabel, "label_2")
        self.label_3 = self.findChild(QLabel, "label_3")
        self.tableWidget = self.findChild(QTableWidget, "tableWidget")
        self.tableWidget_2 = self.findChild(QTableWidget, "tableWidget_2")
        self.pushButton_3 = self.findChild(QPushButton, "pushButton_3")

        for name, w in [('label_2', self.label_2), ('label_3', self.label_3),
                        ('tableWidget', self.tableWidget), ('tableWidget_2', self.tableWidget_2),
                        ('pushButton_3', self.pushButton_3)]:
            if w is None:
                raise RuntimeError(f"В UI отсутствует виджет {name}")

        self.label_2.setText(team1_name)
        self.label_3.setText(team2_name)

    def _setup_tables(self):
        headers = ['Имя', 'Фамилия', 'Дата рождения', 'Позиция']
        for table in [self.tableWidget, self.tableWidget_2]:
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(headers)
            table.setRowCount(10)
            table.setColumnWidth(0, 100)
            table.setColumnWidth(1, 100)
            table.setColumnWidth(2, 120)
            table.setColumnWidth(3, 80)

            for r in range(table.rowCount()):
                for c in range(table.columnCount()):
                    self._ensure_item(table, r, c)

    def _load_existing_players(self):
        try:
            for table, team_id in [(self.tableWidget, self.team1_id), (self.tableWidget_2, self.team2_id)]:
                self.cursor.execute("""
                    SELECT FirstName, LastName, DateOfBirth, Position 
                    FROM Player WHERE ID_Team = %s
                """, (team_id,))
                self._fill_table(table, self.cursor.fetchall())
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить игроков:\n{str(e)}")
            import traceback
            print(traceback.format_exc())

    def _fill_table(self, table, players):
        for r in range(table.rowCount()):
            for c in range(table.columnCount()):
                self._ensure_item(table, r, c).setText("")

        for r, (first_name, last_name, dob, position) in enumerate(players):
            if r >= table.rowCount():
                break
            self._ensure_item(table, r, 0).setText(first_name or "")
            self._ensure_item(table, r, 1).setText(last_name or "")
            self._ensure_item(table, r, 2).setText(dob.strftime('%Y-%m-%d') if dob else "")
            self._ensure_item(table, r, 3).setText(position or "")

    def safe_item_text(self, table, row, col):
        return self._ensure_item(table, row, col).text().strip()

    def _validate_input(self):
        for table in [self.tableWidget, self.tableWidget_2]:
            for row in range(table.rowCount()):
                first_name = self.safe_item_text(table, row, 0)
                last_name = self.safe_item_text(table, row, 1)
                dob_text = self.safe_item_text(table, row, 2)

                if (first_name and not last_name) or (last_name and not first_name):
                    QMessageBox.warning(self, "Ошибка",
                                        f"В строке {row + 1} нужно указать и имя, и фамилию")
                    return False

                if dob_text:
                    qd = QDate.fromString(dob_text, "yyyy-MM-dd")
                    if not qd.isValid():
                        QMessageBox.warning(self, "Ошибка",
                                            f"Некорректная дата в строке {row + 1}. Формат: ГГГГ-ММ-ДД")
                        return False
        return True

    def _save_team_players(self, table, team_id):
        self.cursor.execute("""
            SELECT ID_Player, FirstName, LastName 
            FROM Player WHERE ID_Team = %s
        """, (team_id,))
        existing = {f"{r[1]}_{r[2]}": r[0] for r in self.cursor.fetchall()}

        for row in range(table.rowCount()):
            first_name = self.safe_item_text(table, row, 0)
            last_name = self.safe_item_text(table, row, 1)
            dob_text = self.safe_item_text(table, row, 2)
            position = self.safe_item_text(table, row, 3)

            if not first_name or not last_name:
                continue

            dob = None
            if dob_text:
                qd = QDate.fromString(dob_text, "yyyy-MM-dd")
                if qd.isValid():
                    dob = qd.toPyDate()

            key = f"{first_name}_{last_name}"
            if key in existing:
                self.cursor.execute("""
                    UPDATE Player SET DateOfBirth = %s, Position = %s
                    WHERE ID_Player = %s AND ID_Team = %s
                """, (dob, position, existing[key], team_id))
            else:
                self.cursor.execute("""
                    INSERT INTO Player (FirstName, LastName, DateOfBirth, Position, ID_Team)
                    VALUES (%s, %s, %s, %s, %s)
                """, (first_name, last_name, dob, position, team_id))

    def _save_players(self):
        try:
            if not self._validate_input():
                return

            self._save_team_players(self.tableWidget, self.team1_id)
            self._save_team_players(self.tableWidget_2, self.team2_id)
            
            # Используем новую систему точек сохранения
            transaction_manager = PharmacyTransactionManager(self.conn, self.cursor)
            with transaction_manager.transaction("Сохранение игроков команд") as tx_manager:
                tx_manager.create_savepoint("team1_players_saved")
                tx_manager.create_savepoint("team2_players_saved")
                tx_manager.create_savepoint("players_completed")

            QMessageBox.information(self, "Успех", "Игроки сохранены!")
            self.accept()
        except Exception as e:
            # Используем новую систему отката
            transaction_manager = PharmacyTransactionManager(self.conn, self.cursor)
            transaction_manager.rollback_all()
            import traceback
            QMessageBox.critical(self, "Ошибка",
                                 f"Не удалось сохранить игроков:\n{str(e)}\n\n{traceback.format_exc()}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = LoginWindow()
    window.show()
    sys.exit(app.exec())
